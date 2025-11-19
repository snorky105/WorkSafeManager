import json
import fdb
import bcrypt
import asyncio
from nicegui import ui, app, run
import os
from docx import Document
from datetime import datetime, date
import tempfile
import shutil
import zipfile
import re

# --- LOG DI AVVIO ---
print("--- DEBUG: AVVIO WorkSafeManager (FINAL VERSION - DB CODES) ---")

# --- CONFIGURAZIONE ---
try:
    with open('config.json', 'r') as f:
        config = json.load(f)
except FileNotFoundError:
    print("ATTENZIONE: Config file non trovati. Uso default.")
    config = {'host': 'localhost', 'database': 'database.fdb', 'user': 'SYSDBA', 'password': 'masterkey', 'port': 3050}

# --- DB CONNECTION ---
def get_db_connection():
    return fdb.connect(
        host=config['host'], database=config['database'],
        user=config['user'], password=config['password'],
        port=config.get('port', 3050), charset='UTF8'
    )

# --- HELPERS CALCOLO SESSIONI ---
def get_next_session_number_sync(id_corso, data_svolgimento: date):
    """
    Conta le sessioni (date distinte) per QUESTO specifico corso (id_corso)
    nello stesso mese/anno, precedenti alla data attuale.
    """
    conn = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        
        mese = data_svolgimento.month
        anno = data_svolgimento.year
        
        # Conta le date distinte per QUESTO corso in QUESTO mese
        sql = """
            SELECT COUNT(DISTINCT DATA_SVOLGIMENTO) 
            FROM T_ATTESTATI 
            WHERE ID_CORSO_FK = ? 
            AND EXTRACT(MONTH FROM DATA_SVOLGIMENTO) = ? 
            AND EXTRACT(YEAR FROM DATA_SVOLGIMENTO) = ?
            AND DATA_SVOLGIMENTO < ?
        """
        cur.execute(sql, (id_corso, mese, anno, data_svolgimento))
        row = cur.fetchone()
        count_prev = row[0] if row else 0
        
        return count_prev + 1
        
    except Exception as e:
        print(f"Errore calcolo sessione: {e}")
        return 1 
    finally:
        if conn: conn.close()

# --- REPOSITORY SOGGETTI ---
class UserRepo:
    @staticmethod
    def get_all(search_term=''):
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            sql = "SELECT CODICE_FISCALE, COGNOME, NOME, DATA_NASCITA, LUOGO_NASCITA, ID_ENTE_FK FROM T_SOGGETTI"
            params = []
            if search_term:
                term = search_term.upper()
                sql += " WHERE UPPER(COGNOME) CONTAINING ? OR UPPER(NOME) CONTAINING ? OR UPPER(CODICE_FISCALE) CONTAINING ?"
                params = [term, term, term]
            sql += " ORDER BY COGNOME, NOME"
            cur.execute(sql, tuple(params))
            rows = cur.fetchall()
            result = []
            for r in rows:
                d_nascita = r[3]
                if isinstance(d_nascita, (date, datetime)):
                    d_nascita = d_nascita.strftime('%Y-%m-%d')
                elif d_nascita is None:
                    d_nascita = ''
                result.append({
                    'CODICE_FISCALE': r[0], 'COGNOME': r[1], 'NOME': r[2],
                    'DATA_NASCITA': str(d_nascita), 'LUOGO_NASCITA': r[4], 'ID_ENTE_FK': r[5]
                })
            return result
        except Exception as e:
            print(f"Err UserRepo: {e}")
            return []
        finally:
            if conn: conn.close()

    @staticmethod
    def upsert(data, is_new=True):
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            dt_nascita = None
            if data['DATA_NASCITA']:
                try:
                    if '-' in data['DATA_NASCITA']:
                        dt_nascita = datetime.strptime(data['DATA_NASCITA'], '%Y-%m-%d').date()
                    elif '/' in data['DATA_NASCITA']:
                        dt_nascita = datetime.strptime(data['DATA_NASCITA'], '%d/%m/%Y').date()
                except ValueError: pass
            
            if is_new:
                sql = "INSERT INTO T_SOGGETTI (CODICE_FISCALE, COGNOME, NOME, DATA_NASCITA, LUOGO_NASCITA, ID_ENTE_FK) VALUES (?, ?, ?, ?, ?, ?)"
                params = (data['CODICE_FISCALE'], data['COGNOME'], data['NOME'], dt_nascita, data['LUOGO_NASCITA'], data['ID_ENTE_FK'])
            else:
                sql = "UPDATE T_SOGGETTI SET COGNOME=?, NOME=?, DATA_NASCITA=?, LUOGO_NASCITA=?, ID_ENTE_FK=? WHERE CODICE_FISCALE=?"
                params = (data['COGNOME'], data['NOME'], dt_nascita, data['LUOGO_NASCITA'], data['ID_ENTE_FK'], data['CODICE_FISCALE'])
            cur.execute(sql, params)
            conn.commit()
            return True, "Salvataggio completato."
        except fdb.IntegrityError:
            return False, "Errore: Codice Fiscale esistente o dati invalidi."
        except Exception as e:
            return False, f"Errore DB: {str(e)}"
        finally:
            if conn: conn.close()

    @staticmethod
    def delete(codice_fiscale):
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("DELETE FROM T_SOGGETTI WHERE CODICE_FISCALE = ?", (codice_fiscale,))
            conn.commit()
            return True
        except Exception: return False
        finally:
            if conn: conn.close()

# --- REPOSITORY ENTI ---
class EnteRepo:
    @staticmethod
    def get_all(search_term=''):
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            sql = "SELECT ID_ENTE, DESCRIZIONE, P_IVA FROM T_ENTI"
            params = []
            if search_term:
                term = search_term.upper()
                sql += " WHERE UPPER(DESCRIZIONE) CONTAINING ? OR UPPER(P_IVA) CONTAINING ? OR CAST(ID_ENTE AS VARCHAR(50)) CONTAINING ?"
                params = [term, term, term]
            sql += " ORDER BY DESCRIZIONE"
            cur.execute(sql, tuple(params))
            return [{'ID_ENTE': r[0], 'DESCRIZIONE': r[1], 'P_IVA': r[2]} for r in cur.fetchall()]
        except Exception: return []
        finally:
            if conn: conn.close()

    @staticmethod
    def upsert(data, is_new=True):
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            if is_new:
                sql = "INSERT INTO T_ENTI (ID_ENTE, DESCRIZIONE, P_IVA) VALUES (?, ?, ?)"
                params = (data['ID_ENTE'], data['DESCRIZIONE'], data['P_IVA'])
            else:
                sql = "UPDATE T_ENTI SET DESCRIZIONE=?, P_IVA=? WHERE ID_ENTE=?"
                params = (data['DESCRIZIONE'], data['P_IVA'], data['ID_ENTE'])
            cur.execute(sql, params)
            conn.commit()
            return True, "Salvataggio completato."
        except Exception as e:
            return False, f"Errore DB: {str(e)}"
        finally:
            if conn: conn.close()

    @staticmethod
    def delete(id_ente):
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("DELETE FROM T_ENTI WHERE ID_ENTE = ?", (id_ente,))
            conn.commit()
            return True
        except Exception: return False
        finally:
            if conn: conn.close()

# --- HELPERS RICERCA E DATI ---
def get_user_details_from_db_sync(search_term: str):
    terms = search_term.upper().split()
    if not terms: return []
    sql = "SELECT s.CODICE_FISCALE, s.COGNOME, s.NOME, s.DATA_NASCITA, s.LUOGO_NASCITA, e.DESCRIZIONE FROM T_SOGGETTI s LEFT JOIN T_ENTI e ON s.ID_ENTE_FK = e.ID_ENTE"
    if len(terms) == 1:
        p = terms[0]
        sql += " WHERE (UPPER(s.COGNOME) STARTING WITH ?) OR (UPPER(s.NOME) STARTING WITH ?) OR (UPPER(s.CODICE_FISCALE) STARTING WITH ?)"
        params = [p, p, p]
    else:
        p1, p2 = terms[0], terms[1]
        sql += " WHERE ((UPPER(s.COGNOME) STARTING WITH ? AND UPPER(s.NOME) STARTING WITH ?) OR (UPPER(s.COGNOME) STARTING WITH ? AND UPPER(s.NOME) STARTING WITH ?))"
        params = [p1, p2, p2, p1]
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(sql, tuple(params))
        rows = cur.fetchall()
        conn.close()
        return [{
            "CODICE_FISCALE": r[0], "COGNOME": r[1], "NOME": r[2],
            "DATA_NASCITA": r[3], "LUOGO_NASCITA": r[4], "SOCIETA": r[5] if r[5] else ""
        } for r in rows]
    except Exception: return []

def get_corsi_from_db_sync():
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        # --- MODIFICA: Recuperiamo anche il CODICE_BREVE ---
        cur.execute("SELECT ID_CORSO, NOME_CORSO, ORE_DURATA, CODICE_BREVE FROM T_CORSI ORDER BY NOME_CORSO")
        rows = cur.fetchall()
        conn.close()
        # Restituisce: id, nome, ore, codice
        return [{"id": r[0], "nome": r[1], "ore": r[2], "codice": r[3]} for r in rows]
    except Exception as e:
        print(f"Err Get Corsi: {e}")
        return []

def save_attestato_to_db_sync(cf, id_corso, data_str):
    try:
        dt = None
        if isinstance(data_str, (date, datetime)):
            dt = data_str
        elif re.search(r'\d{4}-\d{2}-\d{2}', str(data_str)):
             dt = datetime.strptime(re.search(r'\d{4}-\d{2}-\d{2}', str(data_str)).group(0), '%Y-%m-%d').date()
        elif re.search(r'\d{2}/\d{2}/\d{4}', str(data_str)):
             dt = datetime.strptime(re.search(r'\d{2}/\d{2}/\d{4}', str(data_str)).group(0), '%d/%m/%Y').date()
        data_val = dt if dt else date.today()
        
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("INSERT INTO T_ATTESTATI (ID_SOGGETTO_FK, ID_CORSO_FK, DATA_SVOLGIMENTO) VALUES (?, ?, ?)", (cf, id_corso, data_val))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        print(f"Err Save Attestato: {e}")
        return False

def generate_certificate_sync(data_map, template_file="modello.docx", output_dir=None):
    if not os.path.exists(template_file): raise FileNotFoundError("Template mancante")
    doc = Document(template_file)
    local_map = data_map.copy()
    dob = local_map.get("{{DATA NASCITA}}")
    if isinstance(dob, (datetime, date)): local_map["{{DATA NASCITA}}"] = dob.strftime('%d/%m/%Y')
    
    # Helper per sostituzione
    def replace_in_p(p, m):
        for k, v in m.items():
            if k in p.text: p.text = p.text.replace(k, str(v if v else ''))

    for p in doc.paragraphs: replace_in_p(p, local_map)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs: replace_in_p(p, local_map)
                         
    fname = f"attestato_{re.sub(r'\W', '', str(local_map.get('{{COGNOME}}','')))}_{re.sub(r'\W', '', str(local_map.get('{{NOME}}','')))}.docx"
    out_path = os.path.join(output_dir, fname) if output_dir else fname
    doc.save(out_path)
    return out_path

def generate_zip_sync(files, base, name="attestati.zip"):
    with zipfile.ZipFile(name, 'w', zipfile.ZIP_DEFLATED) as z:
        for f in files: z.write(f, arcname=os.path.relpath(f, base))
    return name

# --- PAGES ---

@ui.page('/')
def login_page():
    if app.storage.user.get('authenticated', False):
        ui.navigate.to('/dashboard') 
        return
    with ui.column().classes('absolute-center w-full max-w-sm items-center'):
        ui.label("WorkSafeManager").classes("text-3xl font-bold mb-8 text-center")
        with ui.card().style("padding: 40px;").classes("w-full"):
            username_input = ui.input("Utente").props("outlined").classes("w-full mb-2")
            password_input = ui.input("Password").props("outlined type=password").classes("w-full mb-4")
            def on_login_click():
                user = username_input.value.strip()
                pwd = password_input.value.strip()
                if user and pwd: 
                    app.storage.user['authenticated'] = True
                    app.storage.user['username'] = user 
                    ui.notify("Login riuscito!", color="green")
                    ui.navigate.to('/dashboard')
                else:
                    ui.notify("Credenziali errate", color="red")
            ui.button("Entra", on_click=on_login_click).classes("w-full mt-4")

@ui.page('/dashboard')
def dashboard_page():
    if not app.storage.user.get('authenticated', False):
        ui.navigate.to('/') 
        return 
    username = app.storage.user.get('username', 'Utente')
    ui.label(f"Dashboard - {username}").classes("text-3xl font-bold mb-8 text-center")
    with ui.row().classes('w-full justify-center gap-4'):
        with ui.card().classes('w-64 cursor-pointer hover:shadow-xl').on('click', lambda: ui.navigate.to('/creaattestati')):
             ui.icon('explore', size='xl').classes('mx-auto text-primary')
             ui.label('Crea Attestati').classes('text-center text-lg font-bold w-full')
        with ui.card().classes('w-64 cursor-pointer hover:shadow-xl').on('click', lambda: ui.navigate.to('/gestioneutenti')):
             ui.icon('people', size='xl').classes('mx-auto text-primary')
             ui.label('Gestione Utenti').classes('text-center text-lg font-bold w-full')
        with ui.card().classes('w-64 cursor-pointer hover:shadow-xl').on('click', lambda: ui.navigate.to('/gestioneenti')):
             ui.icon('business', size='xl').classes('mx-auto text-primary')
             ui.label('Gestione Enti').classes('text-center text-lg font-bold w-full')
    def logout_click():
        app.storage.user['authenticated'] = False
        ui.navigate.to('/')
    ui.button("Logout", on_click=logout_click).classes("bg-red-500 text-white mx-auto block mt-12")

@ui.page('/creaattestati')
def creaattestati_page():
    if not app.storage.user.get('authenticated', False):
         ui.navigate.to('/')
         return

    # --- CARICAMENTO DATI AGGIORNATO ---
    corsi_raw = get_corsi_from_db_sync()
    
    # Mappa ID -> Nome
    corsi_opts = {c["id"]: c["nome"] for c in corsi_raw}
    # Mappa ID -> Ore
    corsi_ore = {c["id"]: c["ore"] for c in corsi_raw}
    # Mappa ID -> CODICE BREVE DB (Usiamo 'GEN' se vuoto)
    corsi_codici = {c["id"]: (c["codice"].strip() if c["codice"] else "GEN") for c in corsi_raw}
    
    soggetti = {} 

    # Dialogo Ricerca
    search_dialog = ui.dialog()
    with search_dialog, ui.card().classes('w-full max-w-lg'):
        ui.label('Cerca Soggetto').classes('text-xl font-bold mb-2')
        with ui.row().classes('w-full gap-2'):
            search_input = ui.input(label='Cerca...').classes('flex-grow').props('outlined')
            search_btn = ui.button('Cerca').props('color=primary')
        search_results_area = ui.column().classes('w-full mt-2')
        ui.button('Chiudi', on_click=search_dialog.close).props('flat color=grey').classes('ml-auto')

    with ui.column().classes('w-full items-center p-8'):
        with ui.row().classes('w-full items-center mb-4'): 
            ui.button('Torna', on_click=lambda: ui.navigate.to('/dashboard'), icon='arrow_back').props('flat round')
            ui.label('Generazione Attestati Massiva').classes('text-3xl ml-4')
        
        @ui.refreshable
        def render_lista_soggetti():
            if not soggetti:
                ui.label("Nessun soggetto selezionato.").classes('text-sm italic p-4 text-gray-500')
                return
            
            # Header con colonna 'Testo Date Extra'
            with ui.grid().style('grid-template-columns: 1fr 1fr 1fr 2.5fr 1.2fr 1.5fr 0.6fr 0.5fr; width: 100%; font-weight: bold; border-bottom: 2px solid #ccc; padding-bottom: 5px;'):
                ui.label('Cognome'); ui.label('Nome'); ui.label('CF'); ui.label('Corso'); ui.label('Inizio (Data)'); ui.label('Note Date (Es. "e 21/11")'); ui.label('Ore'); ui.label('')

            for cf, item in soggetti.items():
                u_data = item['user']
                if 'date_extra' not in item: item['date_extra'] = ''

                with ui.grid().style('grid-template-columns: 1fr 1fr 1fr 2.5fr 1.2fr 1.5fr 0.6fr 0.5fr; width: 100%; gap: 10px; align-items: center; border-bottom: 1px solid #eee; padding: 5px;'):
                    ui.label(u_data['COGNOME'])
                    ui.label(u_data['NOME'])
                    ui.label(u_data['CODICE_FISCALE']).classes('text-xs')
                    
                    def on_course_change(e, it=item):
                        it['ore'] = corsi_ore.get(e.value)
                    
                    ui.select(options=corsi_opts, on_change=on_course_change).props('outlined dense').bind_value(item, 'cid').classes('w-full')
                    
                    # DATA INIZIO (Usata per conteggio sessione e nome cartella)
                    with ui.input('Data Inizio').props('outlined dense').bind_value(item, 'per').classes('w-full') as date_inp:
                        with date_inp.add_slot('append'):
                            ui.icon('event').classes('cursor-pointer').on('click', lambda: menu.open())
                            with ui.menu() as menu:
                                ui.date().bind_value(item, 'per').on('update:model-value', lambda: menu.close())
                    
                    # TESTO EXTRA (Usato SOLO per stampa su DOCX)
                    ui.input(placeholder='Es: e 21/05').props('outlined dense').bind_value(item, 'date_extra').classes('w-full')
                    
                    ui.number().props('outlined dense').bind_value(item, 'ore')
                    ui.button(icon='delete', on_click=lambda _, c=cf: rimuovi_soggetto(c)).props('flat round dense color=red')

        def process_user_addition(u_data):
            cf = u_data['CODICE_FISCALE']
            if cf in soggetti:
                ui.notify("Presente!", color='orange'); return
            soggetti[cf] = {'user': u_data, 'cid': None, 'per': None, 'date_extra': '', 'ore': None}
            render_lista_soggetti.refresh()
            count_label.set_text(f"Totale: {len(soggetti)}")
            ui.notify(f"Aggiunto: {u_data['COGNOME']} {u_data['NOME']}", color='green')

        def rimuovi_soggetto(cf):
            if cf in soggetti: del soggetti[cf]; render_lista_soggetti.refresh(); count_label.set_text(f"Totale: {len(soggetti)}")
        
        def svuota_lista():
            soggetti.clear(); render_lista_soggetti.refresh(); count_label.set_text("Totale: 0")

        def open_search_ui():
            search_input.value = ""; search_results_area.clear(); search_dialog.open()

        async def perform_search():
            term = search_input.value
            if not term: return
            res = await asyncio.to_thread(get_user_details_from_db_sync, term)
            search_results_area.clear()
            if not res:
                with search_results_area: ui.label("Nessun risultato.").classes('text-red italic')
                return
            if len(res) == 1:
                process_user_addition(res[0]); search_dialog.close()
            else:
                with search_results_area:
                    with ui.list().props('bordered separator dense'):
                        for u in res:
                            dob = u['DATA_NASCITA'].strftime('%d/%m/%Y') if u['DATA_NASCITA'] else "?"
                            lbl = f"{u['COGNOME']} {u['NOME']} ({dob})"
                            with ui.item().props('clickable').on('click', lambda e, x=u: (process_user_addition(x), search_dialog.close())):
                                with ui.item_section():
                                    ui.item_label(lbl)
                                    ui.item_label(u['CODICE_FISCALE']).props('caption')
        search_btn.on_click(perform_search)
        search_input.on('keydown.enter', perform_search)

        with ui.row().classes('w-full justify-between items-center mt-2 mb-2'):
             ui.label('Lista Destinatari').classes('text-xl font-bold')
             with ui.row():
                 ui.button('Aggiungi', on_click=open_search_ui, icon='person_add').props('color=primary')
                 ui.button('Svuota', on_click=svuota_lista, icon='delete_sweep').props('color=red flat')

        with ui.column().classes('w-full p-4 border rounded shadow-md bg-white'):
            count_label = ui.label("Totale: 0").classes('ml-auto text-sm text-gray-500')
            render_lista_soggetti()

        # --- GENERAZIONE PDF/ZIP ---
        async def on_generate():
            items = list(soggetti.values())
            if not items: ui.notify("Lista vuota", color='red'); return
            if any(not x['cid'] or not x['per'] for x in items): ui.notify("Dati mancanti!", color='red'); return

            ui.notify("Generazione...", spinner=True)
            try:
                tmp = tempfile.mkdtemp()
                files_to_zip = []
                
                # 1. Raggruppa per (Corso, DataInizio)
                grouped_items = {}
                for it in items:
                    raw_date = it['per']
                    dt_obj = date.today()
                    try:
                        if re.search(r'\d{4}-\d{2}-\d{2}', str(raw_date)):
                             dt_obj = datetime.strptime(re.search(r'\d{4}-\d{2}-\d{2}', str(raw_date)).group(0), '%Y-%m-%d').date()
                        elif re.search(r'\d{2}/\d{2}/\d{4}', str(raw_date)):
                             dt_obj = datetime.strptime(re.search(r'\d{2}/\d{2}/\d{4}', str(raw_date)).group(0), '%d/%m/%Y').date()
                    except: pass
                    key = (it['cid'], dt_obj)
                    if key not in grouped_items: grouped_items[key] = []
                    grouped_items[key].append(it)

                # 2. Processa Gruppi
                for (cid, dt_val), group_list in grouped_items.items():
                    
                    # A. RECUPERA CODICE DAL DB
                    codice_corso = corsi_codici.get(cid, "GEN")
                    
                    # B. CONTA SESSIONI (Specifico per Corso + Mese)
                    n_sessione = await asyncio.to_thread(get_next_session_number_sync, cid, dt_val)
                    
                    # C. FORMA NOME CARTELLA
                    data_codice = dt_val.strftime('%d%m%Y')
                    sigla_cartella = f"{n_sessione}{codice_corso}{data_codice}"
                    
                    path_sigla = os.path.join(tmp, sigla_cartella)
                    os.makedirs(path_sigla, exist_ok=True)
                    
                    # D. Crea Docx
                    nome_corso_full = corsi_opts[cid]
                    for it in group_list:
                        u = it['user']
                        safe_az = re.sub(r'\W', '_', u.get('SOCIETA', 'Privati'))
                        final_dir = os.path.join(path_sigla, safe_az)
                        os.makedirs(final_dir, exist_ok=True)
                        
                        # Testo Attestato
                        data_inizio_str = dt_val.strftime('%d/%m/%Y')
                        periodo_completo = f"{data_inizio_str} {it['date_extra']}" if it.get('date_extra') else data_inizio_str

                        d_map = {
                            "{{COGNOME}}": u['COGNOME'], "{{NOME}}": u['NOME'],
                            "{{CODICE}}": sigla_cartella,
                            "{{DATA_NASCITA}}": u['DATA_NASCITA'],
                            "{{LUOGO_NASCITA}}": u['LUOGO_NASCITA'],
                            "{{SOCIETA}}": u['SOCIETA'],
                            "{{NOME_CORSO}}": nome_corso_full,
                            "{{DATA_SVOLGIMENTO}}": periodo_completo,
                            "{{ORE_DURATA}}": it['ore'],
                            "{{SIGLA}}": sigla_cartella
                        }
                        f = await asyncio.to_thread(generate_certificate_sync, d_map, "modello.docx", final_dir)
                        files_to_zip.append(f)
                        await asyncio.to_thread(save_attestato_to_db_sync, u['CODICE_FISCALE'], cid, dt_val)

                z_name = f"Export_{datetime.now().strftime('%Y%m%d_%H%M')}.zip"
                z_path = await asyncio.to_thread(generate_zip_sync, files_to_zip, tmp, z_name)
                ui.download(z_path)
                ui.notify(f"Fatto! {len(files_to_zip)} files.", color='green')
                soggetti.clear(); render_lista_soggetti.refresh(); count_label.set_text("Totale: 0")

            except Exception as e:
                ui.notify(f"Errore: {e}", color='red')
                print(f"ERR GEN: {e}")
            finally:
                await asyncio.sleep(10)
                if 'z_path' in locals() and os.path.exists(z_path): os.remove(z_path)
                if os.path.exists(tmp): shutil.rmtree(tmp, ignore_errors=True)

        ui.button("Genera e Scarica", on_click=on_generate).classes('w-full mt-6').props('color=green size=lg')

@ui.page('/gestioneutenti')
def gestioneutenti_page():
    if not app.storage.user.get('authenticated', False): ui.navigate.to('/'); return
    state = {'is_new': True, 'search': ''}
    cf_input = None; cognome_input = None; nome_input = None
    data_input_field = None; luogo_input = None; ente_input = None
    dialog_label = None; table_ref = None; dialog_ref = None

    async def refresh_table():
        rows = await asyncio.to_thread(UserRepo.get_all, state['search'])
        if table_ref: table_ref.rows = rows; table_ref.update()

    def open_dialog(row=None):
        dialog_ref.open()
        if row:
            state['is_new'] = False
            cf_input.value = row['CODICE_FISCALE']; cf_input.props('readonly') 
            cognome_input.value = row['COGNOME']; nome_input.value = row['NOME']
            data_input_field.value = row['DATA_NASCITA']; luogo_input.value = row['LUOGO_NASCITA']
            ente_input.value = row['ID_ENTE_FK']; dialog_label.text = "Modifica Utente"
        else:
            state['is_new'] = True
            cf_input.value = ''; cf_input.props(remove='readonly')
            cognome_input.value = ''; nome_input.value = ''
            data_input_field.value = ''; luogo_input.value = ''; ente_input.value = ''
            dialog_label.text = "Nuovo Utente"

    async def save_user():
        if not cf_input.value or not cognome_input.value: ui.notify('Campi mancanti!', type='warning'); return
        data = {
            'CODICE_FISCALE': cf_input.value.upper().strip(), 'COGNOME': cognome_input.value.strip(),
            'NOME': nome_input.value.strip(), 'DATA_NASCITA': data_input_field.value,
            'LUOGO_NASCITA': luogo_input.value, 'ID_ENTE_FK': ente_input.value
        }
        success, msg = await asyncio.to_thread(UserRepo.upsert, data, state['is_new'])
        if success: ui.notify(msg, type='positive'); dialog_ref.close(); await refresh_table()
        else: ui.notify(msg, type='negative')

    async def delete_user(row):
        await asyncio.to_thread(UserRepo.delete, row['CODICE_FISCALE'])
        ui.notify("Eliminato", type='info'); await refresh_table()

    with ui.column().classes('w-full items-center p-8 max-w-screen-xl mx-auto bg-slate-50 min-h-screen'):
        with ui.row().classes('w-full items-center mb-6 justify-between'):
            ui.button(icon='arrow_back', on_click=lambda: ui.navigate.to('/dashboard')).props('flat round dense')
            ui.label('Utenti').classes('text-3xl font-bold text-slate-800')
            ui.button('Nuovo', icon='add', on_click=lambda: open_dialog(None)).props('unelevated color=primary')

        with ui.card().classes('w-full p-2 mb-4 flex flex-row items-center gap-4'):
            ui.icon('search').classes('text-grey ml-2')
            ui.input(placeholder='Cerca...').classes('flex-grow').props('borderless').bind_value(state, 'search').on('keydown.enter', refresh_table)
            ui.button('Cerca', on_click=refresh_table).props('flat color=primary')

        cols = [
            {'name': 'CODICE_FISCALE', 'label': 'CF', 'field': 'CODICE_FISCALE', 'align': 'left', 'sortable': True},
            {'name': 'COGNOME', 'label': 'Cognome', 'field': 'COGNOME', 'sortable': True, 'align': 'left'},
            {'name': 'NOME', 'label': 'Nome', 'field': 'NOME', 'align': 'left'},
            {'name': 'DATA_NASCITA', 'label': 'Data', 'field': 'DATA_NASCITA', 'align': 'center'},
            {'name': 'azioni', 'label': '', 'field': 'azioni', 'align': 'right'},
        ]
        table_ref = ui.table(columns=cols, rows=[], row_key='CODICE_FISCALE').classes('w-full shadow-md bg-white')
        table_ref.add_slot('body-cell-azioni', r'''
            <q-td key="azioni" :props="props">
                <q-btn icon="edit" size="sm" round flat color="grey-8" @click="$parent.$emit('edit', props.row)" />
                <q-btn icon="delete" size="sm" round flat color="red" @click="$parent.$emit('delete', props.row)" />
            </q-td>
        ''')
        table_ref.on('edit', lambda e: open_dialog(e.args))
        table_ref.on('delete', lambda e: delete_user(e.args))
        ui.timer(0.1, refresh_table, once=True)

    with ui.dialog() as dialog_ref, ui.card().classes('w-full max-w-2xl p-0 rounded-xl overflow-hidden'):
        with ui.row().classes('w-full bg-primary text-white p-4 items-center justify-between'):
            dialog_label = ui.label('Utente').classes('text-lg font-bold')
            ui.button(icon='close', on_click=dialog_ref.close).props('flat round dense text-white')
        with ui.column().classes('w-full p-6 gap-4'):
            with ui.row().classes('w-full gap-4'):
                cf_input = ui.input('CF').props('outlined dense uppercase').classes('w-full md:w-1/2')
                ente_input = ui.input('ID Ente').props('outlined dense').classes('w-full md:w-1/2') 
            with ui.row().classes('w-full gap-4'):
                cognome_input = ui.input('Cognome').props('outlined dense').classes('w-full md:w-1/2')
                nome_input = ui.input('Nome').props('outlined dense').classes('w-full md:w-1/2')
            with ui.row().classes('w-full gap-4'):
                data_input_field = ui.input('Data (YYYY-MM-DD)').props('outlined dense').classes('w-full md:w-1/3')
                luogo_input = ui.input('Luogo Nascita').props('outlined dense').classes('w-full md:w-2/3')
            ui.button('Salva', on_click=save_user).props('unelevated color=primary w-full')

@ui.page('/gestioneenti')
def gestioneenti_page():
    if not app.storage.user.get('authenticated', False): ui.navigate.to('/'); return
    state = {'is_new': True, 'search': ''}
    id_ente_input = None; desc_input = None; piva_input = None
    dialog_label = None; table_ref = None; dialog_ref = None

    async def refresh_table():
        rows = await asyncio.to_thread(EnteRepo.get_all, state['search'])
        if table_ref: table_ref.rows = rows; table_ref.update()

    def open_dialog(row=None):
        dialog_ref.open()
        if row:
            state['is_new'] = False
            id_ente_input.value = row['ID_ENTE']; id_ente_input.props('readonly')
            desc_input.value = row['DESCRIZIONE']; piva_input.value = row['P_IVA']
            dialog_label.text = "Modifica Ente"
        else:
            state['is_new'] = True
            id_ente_input.value = ''; id_ente_input.props(remove='readonly')
            desc_input.value = ''; piva_input.value = ''; dialog_label.text = "Nuovo Ente"

    async def save_ente():
        if not id_ente_input.value or not desc_input.value: ui.notify('Dati mancanti!', type='warning'); return
        data = {'ID_ENTE': id_ente_input.value.strip(), 'DESCRIZIONE': desc_input.value.strip(), 'P_IVA': piva_input.value.strip()}
        success, msg = await asyncio.to_thread(EnteRepo.upsert, data, state['is_new'])
        if success: ui.notify(msg, type='positive'); dialog_ref.close(); await refresh_table()
        else: ui.notify(msg, type='negative')

    async def delete_ente(row):
        await asyncio.to_thread(EnteRepo.delete, row['ID_ENTE'])
        ui.notify("Eliminato", type='info'); await refresh_table()

    with ui.column().classes('w-full items-center p-8 max-w-screen-xl mx-auto bg-slate-50 min-h-screen'):
        with ui.row().classes('w-full items-center mb-6 justify-between'):
            ui.button(icon='arrow_back', on_click=lambda: ui.navigate.to('/dashboard')).props('flat round dense')
            ui.label('Enti').classes('text-3xl font-bold text-slate-800')
            ui.button('Nuovo', icon='add', on_click=lambda: open_dialog(None)).props('unelevated color=primary')

        with ui.card().classes('w-full p-2 mb-4 flex flex-row items-center gap-4'):
            ui.icon('search').classes('text-grey ml-2')
            ui.input(placeholder='Cerca...').classes('flex-grow').props('borderless').bind_value(state, 'search').on('keydown.enter', refresh_table)
            ui.button('Cerca', on_click=refresh_table).props('flat color=primary')

        cols = [
            {'name': 'ID_ENTE', 'label': 'ID', 'field': 'ID_ENTE', 'align': 'left', 'sortable': True},
            {'name': 'DESCRIZIONE', 'label': 'Descrizione', 'field': 'DESCRIZIONE', 'sortable': True, 'align': 'left'},
            {'name': 'P_IVA', 'label': 'P.IVA', 'field': 'P_IVA', 'align': 'center'},
            {'name': 'azioni', 'label': '', 'field': 'azioni', 'align': 'right'},
        ]
        table_ref = ui.table(columns=cols, rows=[], row_key='ID_ENTE').classes('w-full shadow-md bg-white')
        table_ref.add_slot('body-cell-azioni', r'''
            <q-td key="azioni" :props="props">
                <q-btn icon="edit" size="sm" round flat color="grey-8" @click="$parent.$emit('edit', props.row)" />
                <q-btn icon="delete" size="sm" round flat color="red" @click="$parent.$emit('delete', props.row)" />
            </q-td>
        ''')
        table_ref.on('edit', lambda e: open_dialog(e.args))
        table_ref.on('delete', lambda e: delete_ente(e.args))
        ui.timer(0.1, refresh_table, once=True)

    with ui.dialog() as dialog_ref, ui.card().classes('w-full max-w-lg p-0 rounded-xl overflow-hidden'):
        with ui.row().classes('w-full bg-primary text-white p-4 items-center justify-between'):
            dialog_label = ui.label('Ente').classes('text-lg font-bold')
            ui.button(icon='close', on_click=dialog_ref.close).props('flat round dense text-white')
        with ui.column().classes('w-full p-6 gap-4'):
            id_ente_input = ui.input('ID Ente').props('outlined dense').classes('w-full')
            desc_input = ui.input('Ragione Sociale').props('outlined dense').classes('w-full')
            piva_input = ui.input('P.IVA').props('outlined dense').classes('w-full')
            ui.button('Salva', on_click=save_ente).props('unelevated color=primary w-full')

if __name__ in {"__main__", "__mp_main__"}:
    ui.run(title="WorkSafeManager", storage_secret='secret_key', reload=True)