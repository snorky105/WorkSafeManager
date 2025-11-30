import json
import fdb
import psycopg2
import bcrypt
import asyncio
from nicegui import ui, app, run
import os
from docx import Document
from datetime import datetime, date , timedelta
import tempfile
import shutil
import zipfile
import re
import logging
import csv
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

#-- LOGGING --
logging.basicConfig(
    filename='WorkSafeManager.log',
    filemode='a',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger()


# --- LOG DI AVVIO ---
logger.info("WorkSafeManager avviato!")

# --- CONFIGURAZIONE ---
try:
    with open('config_postgres.json', 'r') as f:
        config = json.load(f)
except FileNotFoundError:
    logger.error("ATTENZIONE: File config_postgres.json non trovato. Uso parametri di default.")
    config = {
        'host': 'localhost', 
        'database': 'postgres',  
        'user': 'postgres',
        'password': 'abcd1234', 
        'port': 5432
    }

# --- DB CONNECTION ---
def get_db_connection():
    try:
        return psycopg2.connect(**config)
        
    except psycopg2.Error as e:
        logger.error(f"Errore durante la connessione a PostgreSQL: {e}")
        return None

# -- TEST CONNESSIONE -- 
if __name__ == "__main__":
    conn = get_db_connection()
    if conn:
        logger.info("Connessione a PostgresSQL riuscita")
        conn.close()

# --- HELPERS CALCOLO SESSIONI ---
def get_next_session_number_sync(id_corso, data_svolgimento: date):
    """
    Conta le sessioni (date distinte) per QUESTO specifico corso (id_corso)
    nello stesso mese/anno, precedenti alla data attuale.
    """
    conn = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        
        mese = data_svolgimento.month
        anno = data_svolgimento.year
        
        # Conta le date distinte per QUESTO corso in QUESTO mese
        sql = """
    SELECT COUNT(DISTINCT DATA_SVOLGIMENTO) 
    FROM T_ATTESTATI 
    WHERE ID_CORSO_FK = %s 
    AND EXTRACT(MONTH FROM DATA_SVOLGIMENTO) = %s 
    AND EXTRACT(YEAR FROM DATA_SVOLGIMENTO) = %s
    AND DATA_SVOLGIMENTO < %s
"""

# Nota importante: i parametri devono essere passati nell'ordine esatto dei %s
        cur.execute(sql, (id_corso, mese, anno, data_svolgimento))
        row = cur.fetchone()
        count_prev = row[0] if row else 0
        
        return count_prev + 1
        
    except Exception as e:
        logger.info(f"Errore calcolo sessione: {e}")
        return 1 
    finally:
        if conn: conn.close()

# --- REPOSITORY SOGGETTI ---
class UserRepo:
    @staticmethod
    def get_all(search_term='', solo_docenti=False):
        """
        Recupera utenti. Se solo_docenti=True, filtra solo chi ha IS_DOCENTE=1
        """
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            
            # Selezioniamo anche IS_DOCENTE
            sql = "SELECT CODICE_FISCALE, COGNOME, NOME, DATA_NASCITA, LUOGO_NASCITA, ID_ENTE_FK, IS_DOCENTE FROM T_SOGGETTI"
            
            conditions = []
            params = []
            
            # Filtro Ricerca Testuale
            if search_term:
                term = search_term.upper()
                conditions.append("(UPPER(COGNOME) ILIKE %s OR UPPER(NOME) ILIKE %s OR UPPER(CODICE_FISCALE) ILIKE %s)")
                params.extend([term, term, term])
            
            # Filtro Docente
            if solo_docenti:
                conditions.append("IS_DOCENTE = 1")
            
            # Assembla la query
            if conditions:
                sql += " WHERE " + " AND ".join(conditions)
            
            sql += " ORDER BY COGNOME, NOME"
            
            cur.execute(sql, tuple(params))
            rows = cur.fetchall()
            
            result = []
            for r in rows:
                d_nascita = r[3]
                if isinstance(d_nascita, (date, datetime)):
                    d_nascita = d_nascita.strftime('%Y-%m-%d')
                elif d_nascita is None:
                    d_nascita = ''
                
                result.append({
                    'CODICE_FISCALE': r[0], 
                    'COGNOME': r[1], 
                    'NOME': r[2],
                    'DATA_NASCITA': str(d_nascita), 
                    'LUOGO_NASCITA': r[4], 
                    'ID_ENTE_FK': r[5],
                    'IS_DOCENTE': bool(r[6]) # Convertiamo 0/1 in True/False
                })
            return result
        except Exception as e:
            print(f"Err UserRepo: {e}")
            return []
        finally:
            if conn: conn.close()

    @staticmethod
    def upsert(data, is_new=True):
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            
            # 1. FIX DATA NASCITA
            dt_nascita = None
            raw_data = str(data.get('DATA_NASCITA', '')).strip()
            if raw_data:
                try:
                    if '-' in raw_data:
                        dt_nascita = datetime.strptime(raw_data, '%Y-%m-%d').date()
                    elif '/' in raw_data:
                        dt_nascita = datetime.strptime(raw_data, '%d/%m/%Y').date()
                except ValueError: pass
            
            # 2. FIX ID ENTE (QUESTO È QUELLO CHE MANCAVA)
            # Se è vuoto, lo forziamo a None (NULL su DB)
            id_ente_val = str(data.get('ID_ENTE_FK', '')).strip()
            if not id_ente_val: 
                id_ente_val = None 

            # 3. GESTIONE DOCENTE
            is_doc = 1 if data.get('IS_DOCENTE') else 0

            if is_new:
                sql = """
                    INSERT INTO T_SOGGETTI (CODICE_FISCALE, COGNOME, NOME, DATA_NASCITA, LUOGO_NASCITA, ID_ENTE_FK, IS_DOCENTE) 
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """
                # Nota qui sotto: uso id_ente_val, NON data['ID_ENTE_FK']
                params = (data['CODICE_FISCALE'], data['COGNOME'], data['NOME'], dt_nascita, data['LUOGO_NASCITA'], id_ente_val, is_doc)
            else:
                sql = """
                    UPDATE T_SOGGETTI 
                    SET COGNOME=%s, NOME=%s, DATA_NASCITA=%s, LUOGO_NASCITA=%s, ID_ENTE_FK=%s, IS_DOCENTE=%s 
                    WHERE CODICE_FISCALE=%s
                """
                # Nota qui sotto: uso id_ente_val, NON data['ID_ENTE_FK']
                params = (data['COGNOME'], data['NOME'], dt_nascita, data['LUOGO_NASCITA'], id_ente_val, is_doc, data['CODICE_FISCALE'])
            
            cur.execute(sql, params)
            conn.commit()
            return True, "Salvataggio completato."
        except fdb.IntegrityError:
            return False, "Errore: Codice Fiscale esistente o dati invalidi."
        except Exception as e:
            return False, f"Errore DB: {str(e)}"
        finally:
            if conn: conn.close()
    
    # Delete rimane uguale...
    @staticmethod
    def delete(codice_fiscale):
        # ... (codice uguale a prima) ...
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("DELETE FROM T_SOGGETTI WHERE CODICE_FISCALE = %s", (codice_fiscale,))
            conn.commit()
            return True
        except Exception: return False
        finally:
            if conn: conn.close()

# --- REPOSITORY ATTESTATI ---
class AttestatiRepo:
    @staticmethod
    def get_history(search='', start_date=None, end_date=None):
        """
        Recupera lo storico unendo t_attestati, t_soggetti e t_corsi.
        """
        conn = None
        results = []
        try:
            # Usa la tua funzione di connessione che è già in questo file
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # QUERY CORRETTA PER LA TUA TABELLA t_attestati
            query = """
                SELECT 
                    a.id_attestato,
                    a.data_svolgimento,
                    s.codice_fiscale,
                    s.cognome,
                    s.nome,
                    c.nome_corso
                FROM public.t_attestati a
                JOIN public.t_soggetti s ON a.id_soggetto_fk = s.codice_fiscale
                JOIN public.t_corsi c ON a.id_corso_fk = c.id_corso
                WHERE 1=1
            """
            
            params = []
            
            # Filtro Ricerca
            if search:
                term = f"%{search.lower()}%"
                query += """ AND (
                    LOWER(s.cognome) LIKE %s OR 
                    LOWER(s.nome) LIKE %s OR 
                    LOWER(s.codice_fiscale) LIKE %s OR
                    LOWER(c.nome_corso) LIKE %s
                )"""
                params.extend([term, term, term, term])
            
            # Filtri Date
            if start_date:
                query += " AND a.data_svolgimento >= %s"
                params.append(start_date)
            if end_date:
                query += " AND a.data_svolgimento <= %s"
                params.append(end_date)
                
            query += " ORDER BY a.data_svolgimento DESC"
            
            cursor.execute(query, tuple(params))
            rows = cursor.fetchall()
            
            # Mappatura risultati per la UI
            for row in rows:
                # I dati arrivano nell'ordine della SELECT
                data_svol = row[1]
                
                # Calcolo scadenza (Esempio: 5 anni dopo lo svolgimento)
                # Puoi personalizzarlo o renderlo dipendente dal corso
                try:
                    scadenza = data_svol.replace(year=data_svol.year + 5)
                except:
                    scadenza = None

                results.append({
                    'ID': row[0],
                    'DATA_EMISSIONE': data_svol,
                    'CORSISTA': f"{row[3]} {row[4]}", # Cognome + Nome
                    'CF': row[2],
                    'CORSO': row[5],
                    'SCADENZA': scadenza
                })
                
        except Exception as e:
            print(f"Errore AttestatiRepo: {e}")
        finally:
            if conn: conn.close()
            
        return results
    
    @staticmethod
    def insert_attestato(cf, id_corso, data_svolgimento):
        """
        Salva un nuovo attestato nella tabella t_attestati
        """
        conn = None
        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            
            query = """
                INSERT INTO public.t_attestati 
                (id_soggetto_fk, id_corso_fk, data_svolgimento)
                VALUES (%s, %s, %s)
            """
            cursor.execute(query, (cf, id_corso, data_svolgimento))
            conn.commit()
            return True
        except Exception as e:
            print(f"Errore Insert Attestato: {e}")
            if conn: conn.rollback()
            return False
        finally:
            if conn: conn.close()

# --- REPOSITORY AUTENTICAZIONE ---
class AuthRepo:
    @staticmethod
    def get_all_users():
        """Recupera tutti gli utenti di sistema (senza password hash)"""
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("SELECT USERNAME, RUOLO FROM T_AUTENTICAZIONE ORDER BY USERNAME")
            rows = cur.fetchall()
            return [{'USERNAME': r[0], 'RUOLO': r[1]} for r in rows]
        except Exception as e:
            print(f"Err AuthRepo.get_all: {e}")
            return []
        finally:
            if conn: conn.close()

    @staticmethod
    def create_user(username, password_clear, role='user'):
        """Crea un nuovo utente con password criptata"""
        conn = None
        try:
            # 1. Cripta la password
            pwd_hash = bcrypt.hashpw(password_clear.encode('utf-8'), bcrypt.gensalt())
            
            conn = get_db_connection()
            cur = conn.cursor()
            
            sql = "INSERT INTO T_AUTENTICAZIONE (USERNAME, PASSWORD_HASH, RUOLO) VALUES (%s, %s, %s)"
            cur.execute(sql, (username, pwd_hash.decode('utf-8'), role))
            conn.commit()
            return True, "Utente creato con successo."
        except psycopg2.IntegrityError:
            return False, "Errore: Username già esistente."
            logger.info("Errore: Username già esistente")
        except Exception as e:
            return False, f"Errore DB: {str(e)}"
        finally:
            if conn: conn.close()

    @staticmethod
    def delete_user(username):
        """Elimina un utente"""
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("DELETE FROM T_AUTENTICAZIONE WHERE USERNAME = %s", (username,))
            conn.commit()
            return True
        except Exception: return False
        finally:
            if conn: conn.close()

class CorsoRepo:
    @staticmethod
    def get_all(search=''):
        conn = None
        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # SELECT che include validita_anni
            query = """
                SELECT id_corso, nome_corso, ore_durata, codice_breve, programma, template_file, validita_anni
                FROM public.t_corsi
            """
            params = []
            if search:
                query += " WHERE nome_corso ILIKE %s OR codice_breve ILIKE %s"
                term = f"%{search}%"
                params = [term, term]
            
            query += " ORDER BY id_corso ASC"
            cursor.execute(query, tuple(params))
            
            col_names = [desc[0].upper() for desc in cursor.description]
            results = []
            for row in cursor.fetchall():
                results.append(dict(zip(col_names, row)))
            return results
        except Exception as e:
            print(f"Errore CorsoRepo.get_all: {e}")
            return []
        finally:
            if conn: conn.close()

    @staticmethod
    def get_next_id():
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT COALESCE(MAX(id_corso), 0) + 1 FROM public.t_corsi")
        res = cursor.fetchone()[0]
        conn.close()
        return res

    @staticmethod
    def upsert(data, is_new):
        conn = None
        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # Recuperiamo il valore SENZA default. Se manca è None -> DB NULL
            val_anni = data.get('VALIDITA_ANNI') 

            if is_new:
                # INSERT
                query = """
                    INSERT INTO public.t_corsi 
                    (id_corso, nome_corso, ore_durata, codice_breve, programma, template_file, validita_anni)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """
                cursor.execute(query, (
                    data['ID_CORSO'], data['NOME_CORSO'], data['ORE_DURATA'], 
                    data['CODICE_BREVE'], data['PROGRAMMA'], data['TEMPLATE_FILE'],
                    val_anni # Passa None se vuoto
                ))
                msg = "Corso creato"
            else:
                # UPDATE
                query = """
                    UPDATE public.t_corsi 
                    SET nome_corso=%s, ore_durata=%s, codice_breve=%s, 
                        programma=%s, template_file=%s, validita_anni=%s
                    WHERE id_corso=%s
                """
                cursor.execute(query, (
                    data['NOME_CORSO'], data['ORE_DURATA'], 
                    data['CODICE_BREVE'], data['PROGRAMMA'], data['TEMPLATE_FILE'],
                    val_anni, # Passa None se vuoto
                    data['ID_CORSO']
                ))
                msg = "Corso aggiornato"
            
            conn.commit()
            return True, msg
        except Exception as e:
            if conn: conn.rollback()
            return False, str(e)
        finally:
            if conn: conn.close()

    @staticmethod
    def delete(id_corso):
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM public.t_corsi WHERE id_corso = %s", (id_corso,))
        conn.commit()
        conn.close()
        return True, "Eliminato"

# --- REPOSITORY ENTI (COMPLETA) ---
class EnteRepo:
    @staticmethod
    def get_all(search_term=''):
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            
            sql = "SELECT ID_ENTE, DESCRIZIONE, P_IVA FROM T_ENTI"
            params = []
            
            if search_term:
                term = search_term.upper()
                sql += " WHERE UPPER(DESCRIZIONE) ILIKE %s OR UPPER(P_IVA) ILIKE %s OR CAST(ID_ENTE AS VARCHAR(50)) ILIKE %s"
                params = [term, term, term]
            
            sql += " ORDER BY DESCRIZIONE"
            
            cur.execute(sql, tuple(params))
            rows = cur.fetchall()
            
            # Restituisce la lista di dizionari
            return [{'ID_ENTE': r[0], 'DESCRIZIONE': r[1], 'P_IVA': r[2]} for r in rows]
            
        except Exception as e:
            print(f"Errore EnteRepo.get_all: {e}")
            return []
        finally:
            if conn: conn.close()

    @staticmethod
    def get_next_id():
        """Calcola il prossimo ID disponibile (MAX + 1)"""
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("SELECT MAX(ID_ENTE) FROM T_ENTI")
            row = cur.fetchone()
            max_id = row[0] if row and row[0] is not None else 0
            return max_id + 1
        except Exception as e:
            print(f"Errore calcolo ID Ente: {e}")
            return 1 
        finally:
            if conn: conn.close()

    @staticmethod
    def upsert(data, is_new=True):
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            
            if is_new:
                sql = "INSERT INTO T_ENTI (ID_ENTE, DESCRIZIONE, P_IVA) VALUES (%s, %s, %s)"
                params = (data['ID_ENTE'], data['DESCRIZIONE'], data['P_IVA'])
            else:
                sql = "UPDATE T_ENTI SET DESCRIZIONE=%s, P_IVA=%s WHERE ID_ENTE=%s"
                params = (data['DESCRIZIONE'], data['P_IVA'], data['ID_ENTE'])
            
            cur.execute(sql, params)
            conn.commit()
            return True, "Salvataggio completato."
            
        except Exception as e:
            return False, f"Errore DB: {str(e)}"
        finally:
            if conn: conn.close()

    @staticmethod
    def delete(id_ente):
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("DELETE FROM T_ENTI WHERE ID_ENTE = %s", (id_ente,))
            conn.commit()
            return True
        except Exception:
            return False
        finally:
            if conn: conn.close()

    @staticmethod
    def upsert(data, is_new=True):
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            if is_new:
                sql = "INSERT INTO T_ENTI (ID_ENTE, DESCRIZIONE, P_IVA) VALUES (%s, %s, %s)"
                params = (data['ID_ENTE'], data['DESCRIZIONE'], data['P_IVA'])
            else:
                sql = "UPDATE T_ENTI SET DESCRIZIONE=%s, P_IVA=%s WHERE ID_ENTE=%s"
                params = (data['DESCRIZIONE'], data['P_IVA'], data['ID_ENTE'])
            cur.execute(sql, params)
            conn.commit()
            return True, "Salvataggio completato."
        except Exception as e:
            return False, f"Errore DB: {str(e)}"
        finally:
            if conn: conn.close()

    @staticmethod
    def delete(id_ente):
        conn = None
        try:
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("DELETE FROM T_ENTI WHERE ID_ENTE = %s", (id_ente,))
            conn.commit()
            return True
        except Exception: return False
        finally:
            if conn: conn.close()

    @staticmethod
    def get_history(search='', start_date=None, end_date=None):
        """
        Recupera lo storico unendo t_attestati, t_soggetti e t_corsi.
        """
        conn = None
        results = []
        try:
            # Usa la tua funzione di connessione che è già in questo file
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # QUERY CORRETTA PER LA TUA TABELLA t_attestati
            query = """
                SELECT 
                    a.id_attestato,
                    a.data_svolgimento,
                    s.codice_fiscale,
                    s.cognome,
                    s.nome,
                    c.nome_corso
                FROM public.t_attestati a
                JOIN public.t_soggetti s ON a.id_soggetto_fk = s.codice_fiscale
                JOIN public.t_corsi c ON a.id_corso_fk = c.id_corso
                WHERE 1=1
            """
            
            params = []
            
            # Filtro Ricerca
            if search:
                term = f"%{search.lower()}%"
                query += """ AND (
                    LOWER(s.cognome) LIKE %s OR 
                    LOWER(s.nome) LIKE %s OR 
                    LOWER(s.codice_fiscale) LIKE %s OR
                    LOWER(c.nome_corso) LIKE %s
                )"""
                params.extend([term, term, term, term])
            
            # Filtri Date
            if start_date:
                query += " AND a.data_svolgimento >= %s"
                params.append(start_date)
            if end_date:
                query += " AND a.data_svolgimento <= %s"
                params.append(end_date)
                
            query += " ORDER BY a.data_svolgimento DESC"
            
            cursor.execute(query, tuple(params))
            rows = cursor.fetchall()
            
            # Mappatura risultati per la UI
            for row in rows:
                # I dati arrivano nell'ordine della SELECT
                data_svol = row[1]
                
                # Calcolo scadenza (Esempio: 5 anni dopo lo svolgimento)
                try:
                    scadenza = data_svol.replace(year=data_svol.year + 5)
                except:
                    scadenza = None

                results.append({
                    'ID': row[0],
                    'DATA_EMISSIONE': data_svol,
                    'CORSISTA': f"{row[3]} {row[4]}", # Cognome + Nome
                    'CF': row[2],
                    'CORSO': row[5],
                    'SCADENZA': scadenza
                })
                
        except Exception as e:
            print(f"Errore AttestatiRepo: {e}")
        finally:
            if conn: conn.close()
            
        return results

# --- HELPERS RICERCA E DATI ---
def get_user_details_from_db_sync(search_term: str):
    search_term = search_term.strip()
    if not search_term: return []
    
    # Prepariamo la query base
    sql = """
        SELECT s.CODICE_FISCALE, s.COGNOME, s.NOME, s.DATA_NASCITA, s.LUOGO_NASCITA, e.DESCRIZIONE 
        FROM T_SOGGETTI s 
        LEFT JOIN T_ENTI e ON s.ID_ENTE_FK = e.ID_ENTE
        WHERE 
    """
    
    params = []
    conditions = []

    # 1. CERCA LA FRASE INTERA (Risolve il caso "Di Marco")
    # Cerca se l'intera stringa digitata corrisponde all'inizio di un Cognome, Nome o CF
    full_pattern = search_term + '%'
    conditions.append("(s.COGNOME ILIKE %s OR s.NOME ILIKE %s OR s.CODICE_FISCALE ILIKE %s)")
    params.extend([full_pattern, full_pattern, full_pattern])

    # 2. CERCA LE PAROLE SPEZZATE (Risolve il caso "Rossi Mario")
    # Solo se ci sono almeno due parole separate da spazio
    parts = search_term.split()
    if len(parts) >= 2:
        p1 = parts[0] + '%'
        p2 = parts[1] + '%' 
        # Cerca: (Cognome=Parola1 E Nome=Parola2) OPPURE (Cognome=Parola2 E Nome=Parola1)
        conditions.append("((s.COGNOME ILIKE %s AND s.NOME ILIKE %s) OR (s.COGNOME ILIKE %s AND s.NOME ILIKE %s))")
        params.extend([p1, p2, p2, p1])

    # Unisce le due logiche con "OR" (trova o l'uno o l'altro)
    sql += " (" + " OR ".join(conditions) + ")"
    
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(sql, tuple(params))
        rows = cur.fetchall()
        conn.close()
        return [{
            "CODICE_FISCALE": r[0], "COGNOME": r[1], "NOME": r[2],
            "DATA_NASCITA": r[3], "LUOGO_NASCITA": r[4], "SOCIETA": r[5] if r[5] else ""
        } for r in rows]
    except Exception as e:
        print(f"Errore ricerca: {e}")
        return []

def get_corsi_from_db_sync():
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        
        # Recuperiamo TUTTI i campi, incluso TEMPLATE_FILE (che è il 6° campo, indice 5)
        cur.execute("SELECT ID_CORSO, NOME_CORSO, ORE_DURATA, CODICE_BREVE, PROGRAMMA, TEMPLATE_FILE FROM T_CORSI ORDER BY NOME_CORSO")
        rows = cur.fetchall()
        conn.close()
        
        return [{
            "id": r[0], 
            "nome": r[1], 
            "ore": r[2], 
            "codice": r[3], 
            "programma": r[4] if r[4] else "",
            "template": r[5] if r[5] else "modello.docx" 
        } for r in rows]
        
    except Exception as e:
        print(f"Err Get Corsi: {e}")
        logger.error(f"Errore critico durante la lettura del CorsoRepo: {e}")
        return []

def save_attestato_to_db_sync(cf, id_corso, data_str):
    try:
        dt = None
        if isinstance(data_str, (date, datetime)):
            dt = data_str
        elif re.search(r'\d{4}-\d{2}-\d{2}', str(data_str)):
             dt = datetime.strptime(re.search(r'\d{4}-\d{2}-\d{2}', str(data_str)).group(0), '%Y-%m-%d').date()
        elif re.search(r'\d{2}/\d{2}/\d{4}', str(data_str)):
             dt = datetime.strptime(re.search(r'\d{2}/\d{2}/\d{4}', str(data_str)).group(0), '%d/%m/%Y').date()
        data_val = dt if dt else date.today()
        
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("INSERT INTO T_ATTESTATI (ID_SOGGETTO_FK, ID_CORSO_FK, DATA_SVOLGIMENTO) VALUES (%s, %s, %s)", (cf, id_corso, data_val))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        logger.error(f"Errore durante il salvataggio dell'Attestato: {e}", exc_info=True)
        return False

def get_count_attestati_oggi_sync():
    """Conta gli attestati GENERATI oggi (Data Creazione)"""
    conn = None
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        
        # --- MODIFICA QUI: Usiamo DATA_CREAZIONE invece di DATA_SVOLGIMENTO ---
        cur.execute("SELECT COUNT(*) FROM T_ATTESTATI WHERE DATA_CREAZIONE = CURRENT_DATE")
        
        row = cur.fetchone()
        return row[0] if row else 0
    except Exception as e:
        print(f"Err Count Oggi: {e}")
        return 0
    finally:
        if conn: conn.close()

def generate_certificate_sync(data_map, template_file="modello.docx", output_dir=None):
    if not os.path.exists(template_file): raise FileNotFoundError("Template mancante")
    doc = Document(template_file)
    local_map = data_map.copy()
    
    # --- FIX FORMATO DATA NASCITA ---
    dob = local_map.get("{{DATA_NASCITA}}")
    
    # Caso 1: È un oggetto Data
    if isinstance(dob, (datetime, date)): 
        local_map["{{DATA_NASCITA}}"] = dob.strftime('%d/%m/%Y')
    
    # Caso 2: È una Stringa tipo "1980-05-20"
    elif isinstance(dob, str) and '-' in dob:
        try:
            dt_obj = datetime.strptime(dob.strip(), '%Y-%m-%d')
            local_map["{{DATA_NASCITA}}"] = dt_obj.strftime('%d/%m/%Y')
        except ValueError:
            pass 
    # --------------------------------

    # Helper per sostituzione nel Word
    def replace_in_p(p, m):
        for k, v in m.items():
            if k in p.text: p.text = p.text.replace(k, str(v if v else ''))

    for p in doc.paragraphs: replace_in_p(p, local_map)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs: replace_in_p(p, local_map)
                         
    fname = f"{re.sub(r'\W', '', str(local_map.get('{{COGNOME}}','')))}_{re.sub(r'\W', '', str(local_map.get('{{NOME}}','')))}_{re.sub(r'\W', '', str(local_map.get('{{CODICE}}','')))}_{re.sub(r'\W', '', str(local_map.get('{{SOCIETA}}','')))}.docx"
    out_path = os.path.join(output_dir, fname) if output_dir else fname
    doc.save(out_path)
    return out_path

def generate_zip_sync(files, base, name="attestati.zip"):
    with zipfile.ZipFile(name, 'w', zipfile.ZIP_DEFLATED) as z:
        for f in files: z.write(f, arcname=os.path.relpath(f, base))
    return name

def check_user_credentials_sync(username, plain_password):
    print(f"--- DEBUG LOGIN: Tento accesso per utente '{username}' ---")
    conn = None
    try:
        conn = get_db_connection()
        if not conn:
            print("--- DEBUG: Impossibile connettersi al DB! ---")
            return False
            
        cur = conn.cursor()
        logger.info(" Connessione al DB OK. Cerco utente")
        
        # Cerca l'hash
        cur.execute("SELECT PASSWORD_HASH FROM T_AUTENTICAZIONE WHERE USERNAME = %s", (username,))
        row = cur.fetchone()
        
        if row:
            logger.info("Utente TROVATO nel database")
            stored_hash = row[0]
            
            # Debug dell'hash (stampiamo solo i primi caratteri per sicurezza)
            logger.debug(f"Hash nel DB (primi 10 car): {str(stored_hash)[:10]}...")
        
            # Conversione se necessario
            if isinstance(stored_hash, str):
                stored_hash = stored_hash.encode('utf-8')
            
            # Verifica
            if bcrypt.checkpw(plain_password.encode('utf-8'), stored_hash):
                logger.info("Accesso consentito")
                return True
            else:
                logger.warning("PASSWORD ERRATA! Il controllo bcrypt ha fallito")
        else:
            logger.info(f"Utente : '{username}' NON TROVATO nella tabella T_AUTENTICAZIONE")
                
        return False
        
    except Exception as e:
        logger.warning(" Errore grave!!")
        return False
    finally:
        if conn: conn.close()       

# --- PAGES ---
@ui.page('/')
def login_page():
    # Se l'utente è già loggato, va alla dashboard
    if app.storage.user.get('authenticated', False):
        ui.navigate.to('/dashboard') 
        return
    
    with ui.column().classes('absolute-center w-full max-w-sm items-center'):
        ui.label("WorkSafeManager").classes("text-3xl font-bold mb-8 text-center text-slate-700")
        
        with ui.card().style("padding: 40px;").classes("w-full shadow-xl"):
            username_input = ui.input("Utente").props("outlined").classes("w-full mb-2")
            password_input = ui.input("Password").props("outlined type=password").classes("w-full mb-4")
            
            # --- PUNTO CRITICO: La funzione deve essere async ---
            async def on_login_click():
                user = username_input.value.strip()
                pwd = password_input.value.strip()
                
                # Controllo base: campi vuoti?
                if not user or not pwd:
                    ui.notify("Inserisci utente e password", color="orange")
                    return

                # --- LA VERA VERIFICA DI SICUREZZA ---
                # Chiama la funzione che controlla nel DB (T_AUTENTICAZIONE)
                # Se questa restituisce False, NON entri.
                is_valid = await asyncio.to_thread(check_user_credentials_sync, user, pwd)

                if is_valid:
                    app.storage.user['authenticated'] = True
                    app.storage.user['username'] = user 
                    ui.notify("Login riuscito!", color="green")
                    logger.info("Login riuscito!")
                    ui.navigate.to('/dashboard')
                else:
                    # Se arriva qui, user o password sono sbagliati
                    ui.notify("Credenziali errate o utente non trovato", color="red")
                    logger.error("Credenziali errate o utente non trovato")
            
            # Gestione tasto Invio (per comodità)
            password_input.on('keydown.enter', on_login_click)
            username_input.on('keydown.enter', on_login_click)
            
            ui.button("Entra", on_click=on_login_click).classes("w-full mt-4 bg-primary text-white")

@ui.page('/gestione_accessi')
def gestione_accessi_page():
    # Sicurezza: Se non sei loggato, via!
    if not app.storage.user.get('authenticated', False): 
        ui.navigate.to('/'); return

    # --- 1. INIZIALIZZAZIONE VARIABILI (A NONE) ---
    # Non creiamo ancora gli input, altrimenti appaiono a caso nella pagina
    username_new = None 
    password_new = None
    role_new = None
    
    table_ref = None
    dialog_ref = None

    # --- 2. LOGICA ---
    async def refresh_table():
        rows = await asyncio.to_thread(AuthRepo.get_all_users)
        if table_ref: 
            table_ref.rows = rows
            table_ref.update()

    async def add_new_user():
        # Recuperiamo i valori dagli input che verranno creati dopo
        u = username_new.value.strip()
        p = password_new.value.strip()
        r = role_new.value
        
        if not u or not p:
            ui.notify("Username e Password obbligatori!", type='warning')
            return

        success, msg = await asyncio.to_thread(AuthRepo.create_user, u, p, r)
        if success:
            ui.notify(msg, type='positive')
            dialog_ref.close()
            await refresh_table()
        else:
            ui.notify(msg, type='negative')

    async def delete_user_click(row):
        if row['USERNAME'] == 'admin':
            ui.notify("Non puoi cancellare l'utente admin principale!", type='warning')
            return
            
        await asyncio.to_thread(AuthRepo.delete_user, row['USERNAME'])
        ui.notify(f"Utente {row['USERNAME']} eliminato.", type='info')
        await refresh_table()

    # --- 3. LAYOUT PAGINA ---
    with ui.column().classes('w-full items-center p-8 max-w-screen-lg mx-auto bg-slate-50 min-h-screen'):
        
        # Header
        with ui.row().classes('w-full items-center mb-6 justify-between'):
            ui.button(icon='arrow_back', on_click=lambda: ui.navigate.to('/dashboard')).props('flat round dense')
            ui.label('Gestione Accessi al Software').classes('text-3xl font-bold text-slate-800')
            # Il bottone apre il dialogo (che definiremo sotto)
            ui.button('Nuovo Utente', icon='add', on_click=lambda: dialog_ref.open()).props('unelevated color=primary')

        ui.label("Attenzione: questi sono gli utenti che possono fare LOGIN nel programma.").classes("text-sm text-gray-500 mb-4")

        # Tabella
        cols = [
            {'name': 'USERNAME', 'label': 'Username', 'field': 'USERNAME', 'align': 'left', 'sortable': True},
            {'name': 'RUOLO', 'label': 'Ruolo', 'field': 'RUOLO', 'align': 'center'},
            {'name': 'azioni', 'label': 'Azioni', 'field': 'azioni', 'align': 'right'},
        ]
        
        table_ref = ui.table(columns=cols, rows=[], row_key='USERNAME').classes('w-full shadow-md bg-white')
        
        # Slot Azioni
        table_ref.add_slot('body-cell-azioni', r'''
            <q-td key="azioni" :props="props">
                <q-btn icon="delete" size="sm" round flat color="red" @click="$parent.$emit('delete', props.row)" />
            </q-td>
        ''')
        table_ref.on('delete', lambda e: delete_user_click(e.args))
        
        ui.timer(0.1, refresh_table, once=True)

    # --- 4. DIALOGO (Dove creiamo gli input VERI) ---
    with ui.dialog() as dialog_ref, ui.card().classes('min-w-[400px] p-6'):
        ui.label('Nuovo Utente Sistema').classes('text-xl font-bold mb-4')
        
        # ORA creiamo gli input grafici e li assegniamo alle variabili
        username_new = ui.input('Username').props('outlined').classes('w-full mb-2')
        password_new = ui.input('Password').props('outlined type=password').classes('w-full mb-4')
        role_new = ui.select(options=['admin', 'user', 'segreteria'], label='Ruolo', value='user').props('outlined').classes('w-full mb-6')
        
        with ui.row().classes('w-full justify-end'):
            ui.button('Annulla', on_click=dialog_ref.close).props('flat color=grey')
            ui.button('Crea Utente', on_click=add_new_user).props('unelevated color=primary')

@ui.page('/gestionecorsi')
def gestionecorsi_page():
    if not app.storage.user.get('authenticated', False): 
        ui.navigate.to('/')
        return

    state = {'is_new': True, 'search': '', 'row_to_delete': None}
    
    # --- CONFIGURAZIONE PATH ---
    BASE_DIR = '/home/ubuntu/app/WorkSafeManager' 
    ABSOLUTE_PATH_TO_TEMPLATES = os.path.join(BASE_DIR, 'templates')
    
    print(f"DEBUG: Cartella destinazione template: {ABSOLUTE_PATH_TO_TEMPLATES}")
    
    # Assicurati che la cartella esista, altrimenti creala
    if not os.path.exists(ABSOLUTE_PATH_TO_TEMPLATES):
        try:
            os.makedirs(ABSOLUTE_PATH_TO_TEMPLATES, exist_ok=True)
            print("DEBUG: Cartella templates creata correttamente.")
        except Exception as err:
            ui.notify(f"Errore creazione cartella: {err}", type='negative')
            print(f"ERRORE CRITICO: Non riesco a creare la cartella {ABSOLUTE_PATH_TO_TEMPLATES}. Controlla i permessi!")
    
    # Riferimenti UI
    id_input = None
    nome_input = None
    ore_input = None
    validita_input = None # <--- NUOVO CAMPO VALIDITÀ
    codice_input = None
    programma_input = None
    template_select = None 
    
    dialog_ref = None
    dialog_label = None
    confirm_dialog = None
    table_ref = None

    # --- HELPER ---
    def get_template_files():
        if not os.path.exists(ABSOLUTE_PATH_TO_TEMPLATES):
            try: os.makedirs(ABSOLUTE_PATH_TO_TEMPLATES)
            except: return []
        return [f for f in os.listdir(ABSOLUTE_PATH_TO_TEMPLATES) if f.endswith('.docx') and not f.startswith('~$')]

    async def handle_template_upload(e):
        try:
            # 1. Recupero nome file
            raw_name = getattr(e, 'name', getattr(e, 'filename', None))
            if not raw_name:
                filename = 'modello_caricato.docx'
            else:
                filename = os.path.basename(raw_name)

            # 2. Percorso DIRETTO alla tua cartella specifica
            target_path = os.path.join(ABSOLUTE_PATH_TO_TEMPLATES, filename)
            
            # 3. Scrittura del file
            e.content.seek(0)
            with open(target_path, 'wb') as f:
                f.write(e.content.read())
            
            ui.notify(f"Caricato con successo in: templates/{filename}", type='positive')
            
            # 4. Aggiornamento select
            if template_select:
                template_select.options = get_template_files()
                template_select.value = filename
                template_select.update()
                
        except PermissionError:
            ui.notify("ERRORE PERMESSI: Il server non può scrivere in quella cartella!", type='negative')
        except Exception as ex:
            ui.notify(f"Errore caricamento: {ex}", type='negative')
            
    # --- LOGICA ---
    async def refresh_table():
        rows = await asyncio.to_thread(CorsoRepo.get_all, state['search'])
        if table_ref: 
            table_ref.rows = rows
            table_ref.update()
            
    def open_confirm_delete(row):
        # Salviamo la riga che vogliamo cancellare nello "state"
        state['row_to_delete'] = row
        # Apriamo il dialog che chiede "Sei sicuro?"
        confirm_dialog.open()
        
    async def save_corso():
        # 1. Validazione dati obbligatori
        if not nome_input.value: 
            ui.notify('Il Nome del corso è obbligatorio!', type='warning')
            return
            
        # 2. Conversione sicura dei numeri
        try:
            ore_val = float(ore_input.value) if ore_input.value else 0
            validita_val = int(validita_input.value) if validita_input.value else 5
        except ValueError:
            ui.notify('Errore nei valori numerici (Ore o Validità)', type='warning')
            return
        
        if ore_val <= 0:
            ui.notify('La durata deve essere maggiore di 0', type='warning')
            return

        # 3. Preparazione del dizionario dati
        data = {
            'ID_CORSO': int(id_input.value),
            'NOME_CORSO': nome_input.value.strip(),
            'ORE_DURATA': ore_val,
            'VALIDITA_ANNI': validita_val, 
            'CODICE_BREVE': codice_input.value.strip() if codice_input.value else '',
            'PROGRAMMA': programma_input.value.strip() if programma_input.value else '',
            'TEMPLATE_FILE': template_select.value
        }
        
        # 4. Chiamata al Database (Upsert: Insert o Update)
        # Nota: Assumiamo che CorsoRepo.upsert accetti (dati, is_new)
        success, msg = await asyncio.to_thread(CorsoRepo.upsert, data, state['is_new'])
        
        # 5. Gestione esito
        if success:
            ui.notify(msg, type='positive')
            dialog_ref.close()    # Chiude il popup
            await refresh_table() # Ricarica la tabella
        else:
            ui.notify(f"Errore salvataggio: {msg}", type='negative')

    async def open_dialog(row=None):
        files_disponibili = get_template_files()
        if template_select:
            template_select.options = files_disponibili
            template_select.update()

        dialog_ref.open()
        
        if row:
            state['is_new'] = False
            dialog_label.text = "Modifica Corso"
            id_input.value = row['ID_CORSO']
            nome_input.value = row['NOME_CORSO']
            ore_input.value = row['ORE_DURATA']
            
            # Gestione Validità (Default 5 se manca)
            val_anni = row.get('VALIDITA_ANNI')
            validita_input.value = val_anni if val_anni else 5 
            
            codice_input.value = row['CODICE_BREVE']
            programma_input.value = row['PROGRAMMA']
            
            saved_template = row.get('TEMPLATE_FILE')
            template_select.value = saved_template if saved_template in files_disponibili else None
        else:
            state['is_new'] = True
            dialog_label.text = "Nuovo Corso"
            next_id = await asyncio.to_thread(CorsoRepo.get_next_id)
            id_input.value = next_id
            nome_input.value = ''
            ore_input.value = 8
            validita_input.value = 5 # Default per i nuovi corsi
            codice_input.value = ''
            programma_input.value = ''
            template_select.value = None

    async def execute_delete():
        # 1. Controllo difensivo se lo stato è vuoto
        if not state['row_to_delete']: 
            return

        row = state['row_to_delete']
        
        # 2. Estrazione sicura dell'ID (gestisce sia dict che oggetti)
        if isinstance(row, dict):
            corso_id = row.get('ID_CORSO')
        else:
            corso_id = getattr(row, 'ID_CORSO', None)

        if corso_id is None:
            ui.notify("Errore: ID corso non trovato", type='negative')
            return

        try:
            # 3. Tentativo di cancellazione
            success, msg = await asyncio.to_thread(CorsoRepo.delete, corso_id)
            
            if success:
                ui.notify(msg, type='positive')
                confirm_dialog.close()
                state['row_to_delete'] = None
                await refresh_table()
            else:
                # 4. GESTIONE SPECIFICA ERRORI (Attestati collegati)
                msg_str = str(msg).lower()
                if "foreign key" in msg_str or "t_attestati" in msg_str:
                    ui.notify("Impossibile eliminare: ci sono Attestati emessi per questo corso.", type='warning', timeout=5000)
                    # NON chiudiamo il dialog per dare tempo di leggere
                else:
                    ui.notify(f"Errore DB: {msg}", type='negative')

        except Exception as e:
            # 5. Gestione eccezioni impreviste (es. DB disconnesso)
            err_msg = str(e).lower()
            if "foreign key" in err_msg:
                ui.notify("Impossibile eliminare: ci sono Attestati collegati!", type='warning', timeout=5000)
            else:
                import traceback
                traceback.print_exc()
                ui.notify(f"Errore di sistema: {e}", type='negative')

    # --- LAYOUT ---
    with ui.column().classes('w-full items-center p-8 bg-slate-50 min-h-screen'):
        with ui.row().classes('w-full items-center mb-6 justify-between max-w-screen-xl'):
            with ui.row().classes('items-center gap-4'):
                ui.button(icon='arrow_back', on_click=lambda: ui.navigate.to('/dashboard')).props('flat round dense text-color=slate-700')
                ui.label('Gestione Corsi').classes('text-3xl font-bold text-slate-800')
            ui.button('Nuovo Corso', icon='add', on_click=lambda: open_dialog(None)).props('unelevated color=primary')

        with ui.card().classes('w-full p-2 mb-4 flex flex-row items-center gap-4 max-w-screen-xl'):
            ui.icon('search').classes('text-grey ml-2')
            ui.input(placeholder='Cerca corso...').classes('flex-grow').props('borderless').bind_value(state, 'search').on('keydown.enter', refresh_table)
            ui.button('Cerca', on_click=refresh_table).props('flat color=primary')

        # Tabella
        cols = [
            {'name': 'ID_CORSO', 'label': 'ID', 'field': 'ID_CORSO', 'align': 'left', 'sortable': True, 'style': 'width: 50px'},
            {'name': 'CODICE_BREVE', 'label': 'Codice', 'field': 'CODICE_BREVE', 'align': 'left', 'sortable': True},
            {'name': 'NOME_CORSO', 'label': 'Nome', 'field': 'NOME_CORSO', 'align': 'left', 'sortable': True, 'classes': 'font-bold'},
            {'name': 'ORE_DURATA', 'label': 'Ore', 'field': 'ORE_DURATA', 'align': 'center'},
            {'name': 'VALIDITA_ANNI', 'label': 'Validità (Anni)', 'field': 'VALIDITA_ANNI', 'align': 'center'}, # Nuova Colonna
            {'name': 'TEMPLATE_FILE', 'label': 'Modello', 'field': 'TEMPLATE_FILE', 'align': 'left', 'classes': 'text-gray-500 italic'},
            {'name': 'azioni', 'label': '', 'field': 'azioni', 'align': 'right'},
        ]
        
        table_ref = ui.table(columns=cols, rows=[], row_key='ID_CORSO').classes('w-full shadow-md bg-white max-w-screen-xl')
        
        table_ref.add_slot('body-cell-azioni', r'''
            <q-td key="azioni" :props="props">
                <q-btn icon="edit" size="sm" round flat color="primary" @click="$parent.$emit('edit', props.row)" />
                <q-btn icon="delete" size="sm" round flat color="red" @click="$parent.$emit('delete', props.row)" />
            </q-td>
        ''')
        table_ref.on('edit', lambda e: open_dialog(e.args))
        table_ref.on('delete', lambda e: open_confirm_delete(e.args))
        
        ui.timer(0.1, refresh_table, once=True)

    # --- DIALOG ---
    with ui.dialog() as dialog_ref, ui.card().classes('w-full max-w-2xl p-6 gap-4'):
        dialog_label = ui.label('Corso').classes('text-xl font-bold mb-2 text-slate-800')
        
        with ui.row().classes('w-full gap-4'):
            id_input = ui.input('ID').props('outlined dense readonly').classes('w-20 bg-gray-100')
            nome_input = ui.input('Nome Corso *').props('outlined dense').classes('flex-grow')
            
        with ui.row().classes('w-full gap-4'):
            codice_input = ui.input('Codice Breve').props('outlined dense uppercase').classes('w-1/3')
            
            # ORE e VALIDITA affiancati
            ore_input = ui.number('Ore Durata').props('outlined dense').classes('w-1/4')
            validita_input = ui.number('Validità (Anni)', min=1, max=10).props('outlined dense').classes('w-1/4')

        ui.separator()
        
        with ui.row().classes('w-full items-center gap-2'):
            template_select = ui.select(options=[], label='Seleziona Modello Word (.docx)').props('outlined dense options-dense').classes('flex-grow')
            ui.upload(on_upload=handle_template_upload, auto_upload=True, multiple=False, label='')\
                .props('accept=".docx" flat dense color=primary no-thumbnails icon=cloud_upload').tooltip('Carica Nuovo Modello').classes('w-auto')

        ui.label('Programma Corso').classes('text-sm text-gray-600 font-bold mt-2')
        programma_input = ui.textarea(placeholder='Inserisci dettagli programma...').props('outlined dense rows=6').classes('w-full')

        with ui.row().classes('w-full justify-end mt-4 gap-2'):
            ui.button('Annulla', on_click=dialog_ref.close).props('flat color=grey')
            ui.button('Salva', on_click=save_corso).props('unelevated color=primary icon=save')

    with ui.dialog() as confirm_dialog, ui.card().classes('p-6 items-center text-center'):
        ui.icon('warning', size='xl').classes('text-red-500 mb-2')
        ui.label('Sei sicuro?').classes('text-xl font-bold')
        ui.label('Il corso verrà eliminato.').classes('text-gray-600 mb-4')
        with ui.row().classes('gap-4'):
            ui.button('Annulla', on_click=confirm_dialog.close).props('flat color=grey')
            ui.button('Elimina', on_click=execute_delete).props('unelevated color=red')

@ui.page('/dashboard')
def dashboard_page():
    # Controllo Autenticazione
    if not app.storage.user.get('authenticated', False):
        ui.navigate.to('/') 
        return 
    
    username = app.storage.user.get('username', 'Utente')

    # --- 1. LAYOUT PAGINA ---
    with ui.column().classes('w-full items-center p-8 bg-slate-50 min-h-screen'):
        
        # HEADER
        with ui.row().classes('w-full max-w-5xl justify-between items-center mb-8'):
            with ui.column().classes('gap-0'):
                ui.label(f"Ciao, {username}!").classes("text-3xl font-bold text-slate-800")
                ui.label("Benvenuto nel WorkSafe Manager").classes("text-slate-500")
            
            # Tasto Logout in alto a destra
            def logout_click():
                app.storage.user['authenticated'] = False
                ui.navigate.to('/')
            
            ui.button("Esci", on_click=logout_click, icon='logout').props('flat color=red')

        
        # --- BANNER STATISTICHE ---
        with ui.card().classes('w-full max-w-5xl bg-gradient-to-r from-blue-600 to-indigo-700 text-white p-6 mb-10 shadow-lg rounded-xl flex flex-row items-center justify-between'):
            
            # Parte Sinistra: Testo
            with ui.row().classes('items-center gap-6'):
                with ui.element('div').classes('p-4 bg-white/20 rounded-full'):
                    ui.icon('verified', size='2.5em').classes('text-white')
                
                with ui.column().classes('gap-1'):
                    ui.label('Attestati emessi oggi').classes('text-xl font-semibold')
                    ui.label(datetime.now().strftime('%d %B %Y')).classes('text-sm opacity-80')
            
            # Parte Destra: Il Numero
            with ui.column().classes('items-center bg-white text-blue-800 rounded-lg px-8 py-3 shadow-md min-w-[120px]'):
                count_label = ui.label('...').classes('text-4xl font-bold')
                ui.label('TOTALE').classes('text-[10px] font-bold tracking-widest opacity-60')


        # --- GRIGLIA NAVIGAZIONE ---
        # Usiamo CSS Grid per un allineamento perfetto delle card
        with ui.grid().classes('w-full max-w-5xl grid-cols-1 sm:grid-cols-2 md:grid-cols-3 gap-6'):
            
            def crea_card_menu(titolo, icona, colore, link):
                """Helper per creare le card tutte uguali"""
                color_map = {
                    'indigo': 'text-indigo-600 group-hover:text-indigo-700',
                    'emerald': 'text-emerald-600 group-hover:text-emerald-700',
                    'amber': 'text-amber-600 group-hover:text-amber-700',
                    'rose': 'text-rose-600 group-hover:text-rose-700',
                    'purple': 'text-purple-600 group-hover:text-purple-700',
                    'gray': 'text-gray-600 group-hover:text-gray-800',
                    'cyan': 'text-cyan-600 group-hover:text-cyan-700 bg-cyan-50 group-hover:bg-cyan-100',
                    'orange': 'text-orange-600 group-hover:text-orange-700 bg-orange-50 group-hover:bg-orange-100',
                }
                c_text = color_map.get(colore, 'text-blue-600')

                with ui.card().classes('group cursor-pointer hover:shadow-xl transition-all duration-300 hover:-translate-y-1 h-48 flex flex-col items-center justify-center gap-4 border border-gray-100') \
                     .on('click', lambda: ui.navigate.to(link)):
                     
                     # Icona con cerchio sfondo sfumato
                     with ui.element('div').classes(f'p-4 rounded-full bg-{colore}-50 group-hover:bg-{colore}-100 transition-colors'):
                        ui.icon(icona, size='2.5em').classes(c_text)
                     
                     ui.label(titolo).classes('text-lg font-bold text-slate-700 group-hover:text-slate-900')

            # 1. Crea Attestati (Principale)
            crea_card_menu('Crea Attestati', 'explore', 'indigo', '/creaattestati')
            
            # 2. Gestione Corsi
            crea_card_menu('Gestione Corsi', 'menu_book', 'purple', '/gestionecorsi')
            
            # 3. Gestione Utenti
            crea_card_menu('Anagrafica Utenti', 'people', 'emerald', '/gestioneutenti')
            
            # 4. Gestione Docenti
            crea_card_menu('Gestione Docenti', 'school', 'rose', '/gestionedocenti')
            
            # 5. Gestione Enti
            crea_card_menu('Gestione Enti', 'business', 'amber', '/gestioneenti')
            
            # 6. Gestione Accessi
            crea_card_menu('Gestione Accessi', 'lock_person', 'gray', '/gestione_accessi')
            
            # 7. Archivio
            crea_card_menu('Archivio Storico', 'history', 'cyan', '/archivio')
            
            # 8. Scadenziario
            crea_card_menu('Scadenzario', 'alarm', 'orange', '/scadenzario')


    # --- LOGICA DI AGGIORNAMENTO ---
    async def update_counter():
        # Esegue la funzione DB in un thread separato per non bloccare la UI
        n = await asyncio.to_thread(get_count_attestati_oggi_sync)
        count_label.set_text(f"{n}")
    
    # Aggiorna subito e poi ogni 10 secondi
    ui.timer(0.1, update_counter, once=True) 
    ui.timer(10.0, update_counter)

@ui.page('/creaattestati')
def creaattestati_page():
    if not app.storage.user.get('authenticated', False):
         ui.navigate.to('/')
         return

    # --- CARICAMENTO DATI ---
    corsi_raw = get_corsi_from_db_sync()
    corsi_opts = {c["id"]: c["nome"] for c in corsi_raw}
    corsi_ore = {c["id"]: c["ore"] for c in corsi_raw}
    corsi_codici = {c["id"]: (c["codice"].strip() if c["codice"] else "GEN") for c in corsi_raw}
    
    corsi_programmi = {c["id"]: c["programma"] for c in corsi_raw}
    corsi_templates = {c["id"]: c["template"] for c in corsi_raw}
    
    # --- NOVITÀ: CARICAMENTO DOCENTI ---
    # Recuperiamo solo i docenti per popolare la select
    docenti_list = UserRepo.get_all(solo_docenti=True)
    docenti_opts = {d['CODICE_FISCALE']: f"{d['COGNOME']} {d['NOME']}" for d in docenti_list}

    soggetti = {} 

    # Dialogo Ricerca
    search_dialog = ui.dialog()
    with search_dialog, ui.card().classes('w-full max-w-lg'):
        ui.label('Cerca Soggetto').classes('text-xl font-bold mb-2')
        with ui.row().classes('w-full gap-2'):
            search_input = ui.input(label='Cerca...').classes('flex-grow').props('outlined')
            search_btn = ui.button('Cerca').props('color=primary')
        search_results_area = ui.column().classes('w-full mt-2')
        ui.button('Chiudi', on_click=search_dialog.close).props('flat color=grey').classes('ml-auto')

    with ui.column().classes('w-full items-center p-8'):
        with ui.row().classes('w-full items-center mb-4'): 
            ui.button('Torna', on_click=lambda: ui.navigate.to('/dashboard'), icon='arrow_back').props('flat round')
            ui.label('Generazione Attestati Massiva').classes('text-3xl ml-4')
        
        def estrai_inizio_fine(testo):

            if not testo: return None, None
    
            # Cerca tutte le date nel formato dd/mm/yyyy
            matches = re.findall(r'(\d{2}/\d{2}/\d{4})', str(testo))
            
            valid_dates = []
            for m in matches:
                try:
                    dt = datetime.strptime(m, '%d/%m/%Y').date()
                    valid_dates.append(dt)
                except:
                    pass
            
            if not valid_dates:
                return None, None
            
            # Restituisce (Inizio, Fine)
            return min(valid_dates), max(valid_dates)
        
        @ui.refreshable
        def render_lista_soggetti():
            if not soggetti:
                ui.label("Nessun soggetto selezionato.").classes('text-sm italic p-4 text-gray-500')
                return

            # Definizione colonne griglia
            grid_style = 'grid-template-columns: 0.8fr 0.8fr 0.8fr 2fr 1.2fr 2.5fr 0.4fr 0.3fr; width: 100%; gap: 8px; align-items: start;'
            
            # ---------------------------------------------------------
            # 1. INTESTAZIONE (HEADER) - VA FUORI DAL CICLO FOR!
            # ---------------------------------------------------------
            with ui.grid().style(grid_style + 'font-weight: bold; border-bottom: 2px solid #ccc; padding-bottom: 5px; align-items: center;'):
                ui.label('Cognome')
                ui.label('Nome')
                ui.label('CF')
                
                # Colonna Corso con bottone "Copia Giù" ⬇️
                with ui.row().classes('items-center gap-1'):
                    ui.label('Corso')
                    ui.icon('arrow_downward').classes('cursor-pointer text-blue-400 hover:text-blue-700 text-xs') \
                        .on('click', lambda: applica_a_tutti('cid')).tooltip("Copia il primo su tutti")

                # Colonna Docente con bottone "Copia Giù" ⬇️
                with ui.row().classes('items-center gap-1'):
                    ui.label('Docente')
                    ui.icon('arrow_downward').classes('cursor-pointer text-blue-400 hover:text-blue-700 text-xs') \
                        .on('click', lambda: applica_a_tutti('docente_cf')).tooltip("Copia il primo su tutti")

                # Colonna Calendario con bottone "Copia Giù" ⬇️
                with ui.row().classes('items-center gap-1'):
                    ui.label('Calendario / Orari')
                    ui.icon('arrow_downward').classes('cursor-pointer text-blue-400 hover:text-blue-700 text-xs') \
                        .on('click', lambda: applica_a_tutti('calendario_txt')).tooltip("Copia il primo su tutti")

                ui.label('Ore')
                ui.label('') # Spazio vuoto sopra il cestino

            # ---------------------------------------------------------
            # 2. CORPO DATI (BODY) - VA DENTRO IL CICLO FOR!
            # ---------------------------------------------------------
            for cf, item in soggetti.items():
                u_data = item['user']
                
                # Inizializzazioni di sicurezza
                if 'calendario_txt' not in item: item['calendario_txt'] = ''
                if 'docente_cf' not in item: item['docente_cf'] = None 

                # Inizio riga utente
                with ui.grid().style(grid_style + 'border-bottom: 1px solid #eee; padding: 5px;'):
                    
                    # A. Dati Anagrafici (Solo lettura)
                    ui.label(u_data['COGNOME']).classes('text-sm truncate pt-2 font-medium')
                    ui.label(u_data['NOME']).classes('text-sm truncate pt-2')
                    ui.label(u_data['CODICE_FISCALE']).classes('text-xs truncate pt-2 text-gray-500')
                    
                    # Funzione aggiornamento ore corso
                    def on_course_change(e, it=item):
                        it['ore'] = corsi_ore.get(e.value)
                    
                    # B. Select Corso
                    ui.select(options=corsi_opts, on_change=on_course_change) \
                        .props('outlined dense options-dense') \
                        .bind_value(item, 'cid') \
                        .classes('w-full')
                    
                    # C. Select Docente
                    ui.select(options=docenti_opts) \
                        .props('outlined dense options-dense') \
                        .bind_value(item, 'docente_cf') \
                        .classes('w-full')
                    
                    # D. Textarea Calendario con Feedback (Codice Infallibile)
                    with ui.column().classes('w-full gap-0'):
                        
                        # D1. Textarea (Sopra)
                        t_area = ui.textarea(placeholder='Es: 27/11/2025 14-18...') \
                            .props('outlined dense rows=2 debounce=300') \
                            .bind_value(item, 'calendario_txt') \
                            .classes('w-full text-sm')
                        
                        # D2. Label Feedback (Sotto)
                        lbl_feedback = ui.label('In attesa di inserimento...').classes('text-xs text-gray-400 italic ml-1 mt-1')

                        # D3. Funzione di controllo
                        def aggiorna_live(e):
                            testo = ""
                            if isinstance(e, str): testo = e
                            elif hasattr(e, 'args') and e.args: testo = str(e.args)
                            elif hasattr(e, 'value'): testo = str(e.value)
                            
                            ini, fin = estrai_inizio_fine(testo)
                            
                            if ini and fin:
                                lbl_feedback.text = f"✅ Inizio: {ini.strftime('%d/%m')} - Fine: {fin.strftime('%d/%m')}"
                                lbl_feedback.classes(add='text-green-600 font-bold', remove='text-red-500 text-gray-400')
                            elif len(testo) > 5:
                                lbl_feedback.text = "⚠️ Formato non riconosciuto (usa gg/mm/aaaa)"
                                lbl_feedback.classes(add='text-red-500', remove='text-green-600 text-gray-400 font-bold')
                            else:
                                lbl_feedback.text = "In attesa di date..."
                                lbl_feedback.classes(add='text-gray-400', remove='text-green-600 text-red-500 font-bold')

                        # D4. Collegamento eventi
                        t_area.on('update:model-value', lambda e: aggiorna_live(e))
                        t_area.on('input', lambda e: aggiorna_live(e))
                        
                        # D5. Check iniziale
                        if item['calendario_txt']:
                             aggiorna_live(item['calendario_txt'])
                    
                    # E. Ore
                    ui.number().props('outlined dense').bind_value(item, 'ore').classes('pt-0')
                    
                    # F. Bottone Elimina
                    ui.button(icon='delete', on_click=lambda _, c=cf: rimuovi_soggetto(c)) \
                        .props('flat round dense color=red size=sm').classes('mt-1')

        def process_user_addition(u_data):
            cf = u_data['CODICE_FISCALE']
            if cf in soggetti:
                ui.notify("Presente!", color='orange'); return
            # Inizializziamo anche il docente_cf a None
            soggetti[cf] = {'user': u_data, 'cid': None, 'docente_cf': None, 'per': None, 'date_extra': '', 'ore': None}
            render_lista_soggetti.refresh()
            count_label.set_text(f"Totale: {len(soggetti)}")
            ui.notify(f"Aggiunto: {u_data['COGNOME']}", color='green')
            logger.info(f"Utente: {u_data['COGNOME']} {u_data['NOME']} aggiunto alla griglia!")

        def applica_a_tutti(chiave):

            if not soggetti: 
                ui.notify("Lista vuota, niente da copiare.", color='orange')
                return

            prima_chiave = list(soggetti.keys())[0]
            valore_sorgente = soggetti[prima_chiave].get(chiave)
            
            if not valore_sorgente:
                ui.notify("Il primo rigo è vuoto, impossibile copiare.", color='red')
                return

            # Applica a tutti
            for cf in soggetti:
                soggetti[cf][chiave] = valore_sorgente
                
                # Se stiamo copiando il Corso ('cid'), aggiorniamo anche le ore
                if chiave == 'cid':
                    soggetti[cf]['ore'] = corsi_ore.get(valore_sorgente)

            render_lista_soggetti.refresh()
            ui.notify(f"Dati copiati su {len(soggetti)} righe!", color='positive')
        
        def rimuovi_soggetto(cf):
            if cf in soggetti: del soggetti[cf]; render_lista_soggetti.refresh(); count_label.set_text(f"Totale: {len(soggetti)}")
        
        def svuota_lista():
            soggetti.clear(); render_lista_soggetti.refresh(); count_label.set_text("Totale: 0")

        def open_search_ui():
            search_input.value = ""; search_results_area.clear(); search_dialog.open()

        async def perform_search():
            term = search_input.value
            if not term: return
            res = await asyncio.to_thread(get_user_details_from_db_sync, term)
            search_results_area.clear()
            if not res:
                with search_results_area: ui.label("Nessun risultato.").classes('text-red italic')
                return
            if len(res) == 1:
                process_user_addition(res[0]); search_dialog.close()
            else:
                with search_results_area:
                    with ui.list().props('bordered separator dense'):
                        for u in res:
                            dob = u['DATA_NASCITA'].strftime('%d/%m/%Y') if u['DATA_NASCITA'] else "%s"
                            lbl = f"{u['COGNOME']} {u['NOME']} ({dob})"
                            with ui.item().props('clickable').on('click', lambda e, x=u: (process_user_addition(x), search_dialog.close())):
                                with ui.item_section():
                                    ui.item_label(lbl)
                                    ui.item_label(u['CODICE_FISCALE']).props('caption')
        search_btn.on_click(perform_search)
        search_input.on('keydown.enter', perform_search)

        with ui.row().classes('w-full justify-between items-center mt-2 mb-2'):
             ui.label('Lista Destinatari').classes('text-xl font-bold')
             with ui.row():
                 ui.button('Aggiungi', on_click=open_search_ui, icon='person_add').props('color=primary')
                 ui.button('Svuota', on_click=svuota_lista, icon='delete_sweep').props('color=red flat')

        with ui.column().classes('w-full p-4 border rounded shadow-md bg-white'):
            count_label = ui.label("Totale: 0").classes('ml-auto text-sm text-gray-500')
            render_lista_soggetti()

        # --- GENERAZIONE PDF/ZIP ---
        async def on_generate():
            items = list(soggetti.values())
            if not items: 
                ui.notify("Lista vuota", color='red'); return
            
            # Controllo preliminare: Il corso e il testo calendario devono esserci
            if any(not x['cid'] or not x['calendario_txt'] for x in items): 
                ui.notify("Dati mancanti (Corso o Calendario) per alcuni utenti!", color='red')
                return

            ui.notify("Generazione in corso...", spinner=True)
            
            z_path = None
            tmp = None

            try:
                tmp = tempfile.mkdtemp()
                files_to_zip = []
                
                # Dizionario per raggruppare i file
                grouped_items = {}
                
                # --- FASE 1: ANALISI E RAGGRUPPAMENTO ---
                for it in items:
                    raw_text = it['calendario_txt']
                    
                    # Estraiamo Inizio e Fine dal testo
                    dt_inizio, dt_fine = estrai_inizio_fine(raw_text)
                    
                    if not dt_inizio:
                        # Fallback se l'utente non scrive date valide
                        dt_inizio = date.today()
                        dt_fine = date.today()
                        logger.warning(f"Nessuna data trovata in '{raw_text}', uso data odierna.")
                    
                    # Salviamo le date nell'oggetto per usarle nel prossimo ciclo
                    it['_dt_inizio'] = dt_inizio
                    it['_dt_fine'] = dt_fine

                    # Raggruppiamo basandoci sulla DATA DI INIZIO
                    # (Tutti quelli che iniziano lo stesso giorno nello stesso corso finiscono nella stessa cartella)
                    key = (it['cid'], dt_inizio)
                    if key not in grouped_items: grouped_items[key] = []
                    grouped_items[key].append(it)

                # --- FASE 2: CREAZIONE FILE ---
                for (cid, dt_inizio_val), group_list in grouped_items.items():
                    
                    codice_corso = corsi_codici.get(cid, "GEN")
                    
                    # Calcoliamo il numero sessione basandoci sulla data di INIZIO
                    n_sessione = await asyncio.to_thread(get_next_session_number_sync, cid, dt_inizio_val)
                    data_codice = dt_inizio_val.strftime('%d%m%Y')
                    
                    # Nome cartella: es. 1_SIC_27112025 (Data Inizio)
                    sigla_cartella = f"{n_sessione}{codice_corso}{data_codice}"
                    path_sigla = os.path.join(tmp, sigla_cartella)
                    os.makedirs(path_sigla, exist_ok=True)
                    
                    nome_corso_full = corsi_opts[cid]
                    nome_template = corsi_templates.get(cid, 'modello.docx') or 'modello.docx'
                    programma_txt = corsi_programmi.get(cid, '')
                    path_template = os.path.join("templates", nome_template)
                    
                    if not os.path.exists(path_template):
                        ui.notify(f"ERRORE: Template '{nome_template}' mancante!", color='red')
                        return

                    for it in group_list:
                        u = it['user']
                        safe_az = re.sub(r'\W', '_', u.get('SOCIETA', 'Privati'))
                        final_dir = os.path.join(path_sigla, safe_az)
                        os.makedirs(final_dir, exist_ok=True)
                        
                        nome_docente = docenti_opts.get(it.get('docente_cf'), '')
                        
                        # Testo completo inserito dall'utente (es. "27/11 14-18, 28/11 09-13")
                        periodo_visualizzato = it['calendario_txt']
                        
                        # --- LOGICA RICHIESTA: DATA RILASCIO = ULTIMA DATA ---
                        dt_fine_corrente = it['_dt_fine'] 
                        data_rilascio_str = dt_fine_corrente.strftime('%d/%m/%Y')

                        d_map = {
                            "{{COGNOME}}": u['COGNOME'], 
                            "{{NOME}}": u['NOME'],
                            "{{CODICE}}": sigla_cartella,
                            "{{CF}}": u['CODICE_FISCALE'],
                            "{{DATA_NASCITA}}": u['DATA_NASCITA'],
                            "{{LUOGO_NASCITA}}": u['LUOGO_NASCITA'],
                            "{{SOCIETA}}": u['SOCIETA'],
                            "{{NOME_CORSO}}": nome_corso_full,
                            "{{DATA_SVOLGIMENTO}}": periodo_visualizzato, # Testo libero completo
                            "{{ORE_DURATA}}": it['ore'],
                            "{{DATA_RILASCIOAT}}" : data_rilascio_str,    # Data fine corso
                            "{{SIGLA}}": sigla_cartella,
                            "{{DOCENTE}}": nome_docente,
                            "{{PROGRAMMA}}": programma_txt
                        }

                        f = await asyncio.to_thread(generate_certificate_sync, d_map, path_template, final_dir)
                        files_to_zip.append(f)
                        
                        # Salvataggio DB: Usiamo dt_inizio_val per coerenza con la sessione creata
                        await asyncio.to_thread(save_attestato_to_db_sync, u['CODICE_FISCALE'], cid, dt_inizio_val)

                # --- FASE 3: ZIP E DOWNLOAD ---
                z_name = f"Export_{datetime.now().strftime('%d%m%Y_%H%M%S')}.zip"
                z_path = await asyncio.to_thread(generate_zip_sync, files_to_zip, tmp, z_name)
                
                ui.download(z_path)
                ui.notify(f"Fatto! {len(files_to_zip)} attestati.", color='green')
                
                # Pulizia UI
                soggetti.clear()
                render_lista_soggetti.refresh()
                count_label.set_text("Totale: 0")

            except Exception as e:
                logger.error(f"Errore generazione: {e}")
                ui.notify(f"Errore: {e}", color='red', close_button=True, multi_line=True)
            finally:
                await asyncio.sleep(10)
                if z_path and os.path.exists(z_path): 
                    try: 
                        os.remove(z_path)
                    except: 
                        pass
                if tmp and os.path.exists(tmp): 
                    try: 
                        shutil.rmtree(tmp, ignore_errors=True)
                    except: 
                        pass

        ui.button("Genera attestati", on_click=on_generate).classes('w-full mt-6').props('color=blue size=lg')

@ui.page('/gestioneutenti')
def gestioneutenti_page():
    if not app.storage.user.get('authenticated', False): 
        ui.navigate.to('/')
        return
    
    # Stato locale
    state = {'is_new': True, 'search': '', 'row_to_delete': None}
    
    # Riferimenti UI (Inizializzati a None per sicurezza)
    cf_input = None
    cognome_input = None
    nome_input = None
    data_input_field = None
    luogo_input = None
    ente_select = None
    # societa_input rimosso come richiesto
    is_docente_check = None
    
    dialog_label = None
    table_ref = None
    dialog_ref = None
    confirm_dialog = None 

    # --- HELPER FUNCTIONS ---

    async def get_enti_options():
        """Recupera la lista enti per la select dal DB reale"""
        try:
            enti = await asyncio.to_thread(EnteRepo.get_all, '')
            # Restituisce dict {ID: "Nome (P.IVA)"}
            return {e['ID_ENTE']: f"{e['DESCRIZIONE']} ({e['P_IVA']})" for e in enti}
        except Exception as e:
            ui.notify(f"Errore caricamento Enti: {e}", color='red')
            return {}

    async def refresh_table():
        """Ricarica la tabella dal DB reale"""
        try:
            # 1. Recuperiamo gli utenti
            rows = await asyncio.to_thread(UserRepo.get_all, state['search'])
            
            # 2. Recuperiamo la mappa degli Enti per mostrare il nome nella tabella
            # (Per evitare di mostrare solo l'ID o il campo società vuoto)
            opzioni_enti = await get_enti_options()
            
            # Formattazione dati per visualizzazione
            for r in rows:
                # Formattazione Data (DD-MM-YYYY)
                if r.get('DATA_NASCITA') and '-' in str(r['DATA_NASCITA']):
                    try:
                        d_str = str(r['DATA_NASCITA'])
                        if ' ' in d_str: d_str = d_str.split(' ')[0]
                        anno, mese, giorno = d_str.split('-')[:3]
                        r['DATA_DISPLAY'] = f"{giorno}-{mese}-{anno}"
                    except:
                        r['DATA_DISPLAY'] = r['DATA_NASCITA']
                else:
                    r['DATA_DISPLAY'] = ''

                # Risoluzione Nome Ente da ID
                id_ente = r.get('ID_ENTE_FK')
                # Se l'ID c'è nella mappa usa quello, altrimenti usa stringa vuota o fallback
                if id_ente in opzioni_enti:
                    # opzioni_enti è "Nome (Piva)", prendiamo tutto o solo il nome
                    r['ENTE_DISPLAY'] = opzioni_enti[id_ente]
                else:
                    r['ENTE_DISPLAY'] = '-'

            if table_ref: 
                table_ref.rows = rows
                table_ref.update()
        except Exception as e:
            ui.notify(f"Errore caricamento utenti: {e}", color='red')

    async def open_dialog(row=None):
        """Apre il dialog Creazione/Modifica"""
        # Carica opzioni Enti aggiornate
        opzioni_enti = await get_enti_options()
        if ente_select:
            ente_select.options = opzioni_enti
            ente_select.update()

        dialog_ref.open()
        
        if row:
            # --- MODIFICA ---
            state['is_new'] = False
            dialog_label.text = "Modifica Utente"
            
            # Popolamento campi
            cf_input.value = row['CODICE_FISCALE']
            cf_input.props('readonly') # CF bloccato in modifica
            
            cognome_input.value = row['COGNOME']
            nome_input.value = row['NOME']
            
            data_val = str(row['DATA_NASCITA']).split(' ')[0] if row['DATA_NASCITA'] else None
            data_input_field.value = data_val
            
            luogo_input.value = row['LUOGO_NASCITA']
            is_docente_check.value = bool(row.get('IS_DOCENTE', False))
            
            # Gestione Ente
            current_ente = row.get('ID_ENTE_FK')
            ente_select.value = current_ente if current_ente in opzioni_enti else None
            
        else:
            # --- NUOVO ---
            state['is_new'] = True
            dialog_label.text = "Nuovo Utente"
            
            # Reset campi
            cf_input.value = ''
            cf_input.props(remove='readonly') # CF sbloccato
            cognome_input.value = ''
            nome_input.value = ''
            data_input_field.value = None
            luogo_input.value = ''
            is_docente_check.value = False
            ente_select.value = None

    async def save_user():
        """Salva nel DB"""
        if not cf_input.value or not cognome_input.value: 
            ui.notify('Codice Fiscale e Cognome obbligatori!', type='warning')
            return
        
        ente_val = ente_select.value if ente_select.value is not None else None

        data = {
            'CODICE_FISCALE': cf_input.value.upper().strip(), 
            'COGNOME': cognome_input.value.strip(),
            'NOME': nome_input.value.strip(), 
            'DATA_NASCITA': data_input_field.value,
            'LUOGO_NASCITA': luogo_input.value.strip() if luogo_input.value else '', 
            'IS_DOCENTE': is_docente_check.value,
            'ID_ENTE_FK': ente_val,
            # Campo SOCIETA rimosso o impostato vuoto se il DB lo richiede ancora
            'SOCIETA': '' 
        }
        
        try:
            success, msg = await asyncio.to_thread(UserRepo.upsert, data, state['is_new'])
            
            if success: 
                ui.notify(msg, type='positive')
                dialog_ref.close()
                await refresh_table()
            else: 
                ui.notify(msg, type='negative')
        except Exception as e:
            ui.notify(f"Errore salvataggio: {e}", color='red')

    # --- LOGICA ELIMINAZIONE ---
    def open_confirm_delete(row):
        state['row_to_delete'] = row
        confirm_dialog.open()

    async def execute_delete():
        if not state['row_to_delete']: return
        row = state['row_to_delete']
        
        try:
            success, msg = await asyncio.to_thread(UserRepo.delete, row['CODICE_FISCALE'])
            
            if success:
                ui.notify("Utente eliminato", type='positive')
                await refresh_table()
            else:
                ui.notify(f"Errore: {msg}", type='negative')
        except Exception as e:
            ui.notify(f"Errore eliminazione: {e}", color='red')
            
        confirm_dialog.close()
        state['row_to_delete'] = None

    # --- LAYOUT PAGINA ---
    with ui.column().classes('w-full items-center p-8 max-w-screen-xl mx-auto bg-slate-50 min-h-screen'):
        
        # Header
        with ui.row().classes('w-full items-center mb-6 justify-between'):
            with ui.row().classes('items-center gap-4'):
                ui.button(icon='arrow_back', on_click=lambda: ui.navigate.to('/dashboard')).props('flat round dense text-color=slate-700')
                ui.label('Gestione Utenti').classes('text-3xl font-bold text-slate-800')
            
            ui.button('Nuovo Utente', icon='person_add', on_click=lambda: open_dialog(None)).props('unelevated color=primary')

        # Barra Ricerca
        with ui.card().classes('w-full p-2 mb-4 flex flex-row items-center gap-4'):
            ui.icon('search').classes('text-grey ml-2')
            ui.input(placeholder='Cerca per cognome o CF...').classes('flex-grow').props('borderless').bind_value(state, 'search').on('keydown.enter', refresh_table)
            ui.button('Cerca', on_click=refresh_table).props('flat color=primary')

        # Tabella
        cols = [
            {'name': 'CODICE_FISCALE', 'label': 'CF', 'field': 'CODICE_FISCALE', 'align': 'left', 'sortable': True, 'classes': 'font-mono text-xs'},
            {'name': 'COGNOME', 'label': 'Cognome', 'field': 'COGNOME', 'sortable': True, 'align': 'left', 'classes': 'font-bold'},
            {'name': 'NOME', 'label': 'Nome', 'field': 'NOME', 'align': 'left', 'classes': 'font-bold'},
            {'name': 'DATA_NASCITA', 'label': 'Data', 'field': 'DATA_DISPLAY', 'align': 'center'}, 
            # Modificato per mostrare il nome dell'Ente invece del campo Società libero
            {'name': 'ENTE_DISPLAY', 'label': 'Ente / Azienda', 'field': 'ENTE_DISPLAY', 'align': 'left'},
            {'name': 'IS_DOCENTE', 'label': 'Ruolo', 'field': 'IS_DOCENTE', 'align': 'center'},
            {'name': 'azioni', 'label': '', 'field': 'azioni', 'align': 'right'},
        ]
        
        table_ref = ui.table(columns=cols, rows=[], row_key='CODICE_FISCALE').classes('w-full shadow-md bg-white')
        
        # Slot Badge Docente
        table_ref.add_slot('body-cell-IS_DOCENTE', r'''
            <q-td key="IS_DOCENTE" :props="props">
                <q-badge :color="props.value ? 'purple' : 'grey'" :label="props.value ? 'DOCENTE' : 'CORSISTA'" outline />
            </q-td>
        ''')

        # Slot Azioni
        table_ref.add_slot('body-cell-azioni', r'''
            <q-td key="azioni" :props="props">
                <q-btn icon="edit" size="sm" round flat color="grey-8" @click="$parent.$emit('edit', props.row)" />
                <q-btn icon="delete" size="sm" round flat color="red" @click="$parent.$emit('delete', props.row)" />
            </q-td>
        ''')
        
        table_ref.on('edit', lambda e: open_dialog(e.args))
        table_ref.on('delete', lambda e: open_confirm_delete(e.args))
        
        ui.timer(0.1, refresh_table, once=True)

    # --- DIALOG EDIT/NEW ---
    with ui.dialog() as dialog_ref, ui.card().classes('w-full max-w-2xl p-0 rounded-xl overflow-hidden'):
        
        # Header Dialogo
        with ui.row().classes('w-full bg-primary text-white p-4 items-center justify-between'):
            dialog_label = ui.label('Utente').classes('text-lg font-bold')
            ui.button(icon='close', on_click=dialog_ref.close).props('flat round dense text-white')
        
        with ui.column().classes('w-full p-6 gap-4'):
            
            # Riga 1: CF e Docente
            with ui.row().classes('w-full gap-4 items-center'):
                cf_input = ui.input('Codice Fiscale *').props('outlined dense uppercase maxlength=16').classes('w-full md:w-1/2')
                is_docente_check = ui.checkbox('Abilita come Docente').classes('text-purple-600 font-bold')

            # Riga 2: Anagrafica
            with ui.row().classes('w-full gap-4'):
                cognome_input = ui.input('Cognome *').props('outlined dense').classes('w-full md:w-1/2')
                nome_input = ui.input('Nome *').props('outlined dense').classes('w-full md:w-1/2')
            
            # Riga 3: Dati Nascita
            with ui.row().classes('w-full gap-4'):
                data_input_field = ui.input('Data (YYYY-MM-DD)').props('outlined dense type=date').classes('w-full md:w-1/3')
                luogo_input = ui.input('Luogo Nascita').props('outlined dense').classes('w-full md:w-2/3')
            
            # Riga 4: Solo Ente (Rimosso Società input libero)
            ente_select = ui.select(options={}, with_input=True, label='Ente di Appartenenza') \
                .props('outlined dense use-input behavior="menu" clearable') \
                .classes('w-full')

            ui.separator().classes('mt-4')
            ui.button('Salva Utente', on_click=save_user).props('unelevated color=primary w-full icon=save')

    # --- DIALOG CONFERMA ELIMINAZIONE ---
    with ui.dialog() as confirm_dialog, ui.card().classes('p-6 items-center text-center'):
        ui.icon('warning', size='xl').classes('text-red-500 mb-2')
        ui.label('Confermi eliminazione?').classes('text-xl font-bold text-slate-800')
        ui.label('L\'utente verrà rimosso permanentemente.').classes('text-gray-600 mb-4')
        with ui.row().classes('gap-4'):
            ui.button('Annulla', on_click=confirm_dialog.close).props('flat color=grey')
            ui.button('Elimina', on_click=execute_delete).props('unelevated color=red')

@ui.page('/gestioneenti')
def gestioneenti_page():
    if not app.storage.user.get('authenticated', False): ui.navigate.to('/'); return
    
    state = {'is_new': True, 'search': ''}
    
    # --- 1. INIZIALIZZIAMO LE VARIABILI QUI ---
    # Questo è fondamentale affinché open_dialog le veda, anche se sono vuote all'inizio
    # Usiamo un "contenitore" (lista o dict) o semplicemente dichiariamo le variabili
    # Ma il trucco migliore con NiceGUI è definire i riferimenti prima
    
    dialog_ref = None
    dialog_label = None
    table_ref = None
    
    # Definiamo gli input ma li creeremo dopo nella UI. 
    # Li inizializziamo a None per evitare il NameError, ma Python deve sapere che esistono.
    id_ente_input = None 
    desc_input = None 
    piva_input = None

    # --- 2. LOGICA ---

    async def refresh_table():
        rows = await asyncio.to_thread(EnteRepo.get_all, state['search'])
        if table_ref: table_ref.rows = rows; table_ref.update()

    async def open_dialog(row=None):
        # NOTA: Qui usiamo 'nonlocal' se dovessimo riassegnare l'oggetto input intero, 
        # ma dato che tocchiamo solo .value e .props, Python capisce il riferimento 
        # SE l'oggetto è stato creato prima della chiamata della funzione.
        
        dialog_ref.open()
        if row:
            # --- MODALITÀ MODIFICA ---
            state['is_new'] = False
            id_ente_input.value = row['ID_ENTE']
            id_ente_input.props('readonly') # Blocchiamo ID
            desc_input.value = row['DESCRIZIONE']
            piva_input.value = row['P_IVA']
            dialog_label.text = "Modifica Ente"
        else:
            # --- MODALITÀ NUOVO (AUTO-INCREMENT) ---
            state['is_new'] = True
            
            # Calcolo ID automatico
            next_id = await asyncio.to_thread(EnteRepo.get_next_id)
            
            id_ente_input.value = next_id     # Precompila
            id_ente_input.props('readonly')   # Blocca modifica
            
            desc_input.value = ''
            piva_input.value = ''
            dialog_label.text = "Nuovo Ente"

    async def save_ente():
        if not id_ente_input.value or not desc_input.value: 
            ui.notify('Dati mancanti!', type='warning')
            return
        
        id_val = str(id_ente_input.value).strip()
        desc_val = str(desc_input.value).strip()
        piva_val = str(piva_input.value).strip() if piva_input.value else ''

        data = {'ID_ENTE': id_val, 'DESCRIZIONE': desc_val, 'P_IVA': piva_val}

        success, msg = await asyncio.to_thread(EnteRepo.upsert, data, state['is_new'])
        if success: 
            ui.notify(msg, type='positive')
            dialog_ref.close()
            await refresh_table()
        else: 
            ui.notify(msg, type='negative')

    async def delete_ente(row):
        await asyncio.to_thread(EnteRepo.delete, row['ID_ENTE'])
        ui.notify("Eliminato", type='info'); await refresh_table()

    # --- 3. INTERFACCIA UTENTE (UI) ---
    
    with ui.column().classes('w-full items-center p-8 max-w-screen-xl mx-auto bg-slate-50 min-h-screen'):
        with ui.row().classes('w-full items-center mb-6 justify-between'):
            ui.button(icon='arrow_back', on_click=lambda: ui.navigate.to('/dashboard')).props('flat round dense')
            ui.label('Enti').classes('text-3xl font-bold text-slate-800')
            # NOTA: Qui passiamo open_dialog, che ora è async ma NiceGUI gestisce lambda/async bene
            ui.button('Nuovo', icon='add', on_click=lambda: open_dialog(None)).props('unelevated color=primary')

        with ui.card().classes('w-full p-2 mb-4 flex flex-row items-center gap-4'):
            ui.icon('search').classes('text-grey ml-2')
            ui.input(placeholder='Cerca...').classes('flex-grow').props('borderless').bind_value(state, 'search').on('keydown.enter', refresh_table)
            ui.button('Cerca', on_click=refresh_table).props('flat color=primary')

        cols = [
            {'name': 'ID_ENTE', 'label': 'ID', 'field': 'ID_ENTE', 'align': 'left', 'sortable': True},
            {'name': 'DESCRIZIONE', 'label': 'Descrizione', 'field': 'DESCRIZIONE', 'sortable': True, 'align': 'left'},
            {'name': 'P_IVA', 'label': 'P.IVA', 'field': 'P_IVA', 'align': 'center'},
            {'name': 'azioni', 'label': '', 'field': 'azioni', 'align': 'right'},
        ]
        table_ref = ui.table(columns=cols, rows=[], row_key='ID_ENTE').classes('w-full shadow-md bg-white')
        table_ref.add_slot('body-cell-azioni', r'''
            <q-td key="azioni" :props="props">
                <q-btn icon="edit" size="sm" round flat color="grey-8" @click="$parent.$emit('edit', props.row)" />
                <q-btn icon="delete" size="sm" round flat color="red" @click="$parent.$emit('delete', props.row)" />
            </q-td>
        ''')
        table_ref.on('edit', lambda e: open_dialog(e.args))
        table_ref.on('delete', lambda e: delete_ente(e.args))
        ui.timer(0.1, refresh_table, once=True)

    # --- 4. CREAZIONE DIALOG E INPUT ---
    # Qui vengono effettivamente create le variabili id_ente_input, ecc.
    
    with ui.dialog() as dialog_ref, ui.card().classes('w-full max-w-lg p-0 rounded-xl overflow-hidden'):
        with ui.row().classes('w-full bg-primary text-white p-4 items-center justify-between'):
            dialog_label = ui.label('Ente').classes('text-lg font-bold')
            ui.button(icon='close', on_click=dialog_ref.close).props('flat round dense text-white')
        
        with ui.column().classes('w-full p-6 gap-4'):
            # --- PUNTO CRUCIALE: Assegniamo alle variabili inizializzate all'inizio ---
            id_ente_input = ui.input('ID Ente').props('outlined dense').classes('w-full')
            desc_input = ui.input('Ragione Sociale').props('outlined dense').classes('w-full')
            piva_input = ui.input('P.IVA').props('outlined dense').classes('w-full')
            
            ui.button('Salva', on_click=save_ente).props('unelevated color=primary w-full')
            
@ui.page('/gestionedocenti')
def gestionedocenti_page():
    # -- PROVA -- 
    if not app.storage.user.get('authenticated', False): ui.navigate.to('/'); return
    state = {'is_new': True, 'search': ''}
    # Variabili UI
    cf_input = None; cognome_input = None; nome_input = None
    data_input_field = None; luogo_input = None; ente_input = None
    dialog_ref = None; table_ref = None; dialog_label = None

    async def refresh_table():
        # --- QUI LA DIFFERENZA: solo_docenti=True ---
        rows = await asyncio.to_thread(UserRepo.get_all, state['search'], solo_docenti=True)
        if table_ref: table_ref.rows = rows; table_ref.update()

    def open_dialog(row=None):
        dialog_ref.open()
        if row:
            state['is_new'] = False
            cf_input.value = row['CODICE_FISCALE']; cf_input.props('readonly') 
            cognome_input.value = row['COGNOME']; nome_input.value = row['NOME']
            data_input_field.value = row['DATA_NASCITA']; luogo_input.value = row['LUOGO_NASCITA']
            ente_input.value = row['ID_ENTE_FK']
            dialog_label.text = "Modifica Docente"
        else:
            state['is_new'] = True
            cf_input.value = ''; cf_input.props(remove='readonly')
            cognome_input.value = ''; nome_input.value = ''
            data_input_field.value = ''; luogo_input.value = ''; ente_input.value = ''
            dialog_label.text = "Nuovo Docente"

    async def save_docente():
        if not cf_input.value or not cognome_input.value: ui.notify('Dati mancanti!', type='warning'); return
        
        data = {
            'CODICE_FISCALE': cf_input.value.upper().strip(),
            'COGNOME': cognome_input.value.strip(),
            'NOME': nome_input.value.strip(),
            'DATA_NASCITA': data_input_field.value,
            'LUOGO_NASCITA': luogo_input.value,
            'ID_ENTE_FK': ente_input.value,
            'IS_DOCENTE': True  # --- FORZIAMO CHE SIA UN DOCENTE ---
        }
        
        success, msg = await asyncio.to_thread(UserRepo.upsert, data, state['is_new'])
        if success: ui.notify(msg, type='positive'); dialog_ref.close(); await refresh_table()
        else: ui.notify(msg, type='negative')

    async def delete_docente(row):
        await asyncio.to_thread(UserRepo.delete, row['CODICE_FISCALE'])
        ui.notify("Eliminato", type='info'); await refresh_table()

    # --- UI IDENTICA A GESTIONE UTENTI MA TITOLI DIVERSI ---
    with ui.column().classes('w-full items-center p-8 max-w-screen-xl mx-auto bg-slate-50 min-h-screen'):
        with ui.row().classes('w-full items-center mb-6 justify-between'):
            ui.button(icon='arrow_back', on_click=lambda: ui.navigate.to('/dashboard')).props('flat round dense')
            ui.label('Gestione Docenti').classes('text-3xl font-bold text-slate-800')
            ui.button('Nuovo Docente', icon='add', on_click=lambda: open_dialog(None)).props('unelevated color=primary')

        # ... SEARCH BAR ...
        with ui.card().classes('w-full p-2 mb-4 flex flex-row items-center gap-4'):
            ui.icon('search').classes('text-grey ml-2')
            ui.input(placeholder='Cerca Docente...').classes('flex-grow').props('borderless').bind_value(state, 'search').on('keydown.enter', refresh_table)
            ui.button('Cerca', on_click=refresh_table).props('flat color=primary')

        # ... TABLE ...
        cols = [
            {'name': 'CODICE_FISCALE', 'label': 'CF', 'field': 'CODICE_FISCALE', 'align': 'left', 'sortable': True},
            {'name': 'COGNOME', 'label': 'Cognome', 'field': 'COGNOME', 'sortable': True, 'align': 'left'},
            {'name': 'NOME', 'label': 'Nome', 'field': 'NOME', 'align': 'left'},
            {'name': 'azioni', 'label': '', 'field': 'azioni', 'align': 'right'},
        ]
        table_ref = ui.table(columns=cols, rows=[], row_key='CODICE_FISCALE').classes('w-full shadow-md bg-white')
        table_ref.add_slot('body-cell-azioni', r'''
            <q-td key="azioni" :props="props">
                <q-btn icon="edit" size="sm" round flat color="grey-8" @click="$parent.$emit('edit', props.row)" />
                <q-btn icon="delete" size="sm" round flat color="red" @click="$parent.$emit('delete', props.row)" />
            </q-td>
        ''')
        table_ref.on('edit', lambda e: open_dialog(e.args))
        table_ref.on('delete', lambda e: delete_docente(e.args))
        ui.timer(0.1, refresh_table, once=True)

    # --- DIALOGO (Identico a Gestione Utenti) ---
    with ui.dialog() as dialog_ref, ui.card().classes('w-full max-w-2xl p-0 rounded-xl overflow-hidden'):
        with ui.row().classes('w-full bg-primary text-white p-4 items-center justify-between'):
            dialog_label = ui.label('Docente').classes('text-lg font-bold')
            ui.button(icon='close', on_click=dialog_ref.close).props('flat round dense text-white')
        with ui.column().classes('w-full p-6 gap-4'):
            with ui.row().classes('w-full gap-4'):
                cf_input = ui.input('Codice Fiscale *').props('outlined dense uppercase maxlength=16').classes('w-full md:w-1/2')
                ente_input = ui.input('ID Ente (Opzionale)').props('outlined dense').classes('w-full md:w-1/2') 
            with ui.row().classes('w-full gap-4'):
                cognome_input = ui.input('Cognome').props('outlined dense').classes('w-full md:w-1/2')
                nome_input = ui.input('Nome').props('outlined dense').classes('w-full md:w-1/2')
            with ui.row().classes('w-full gap-4'):
                data_input_field = ui.input('Data Nascita').props('outlined dense').classes('w-full md:w-1/3')
                luogo_input = ui.input('Luogo Nascita').props('outlined dense').classes('w-full md:w-2/3')
            ui.button('Salva Docente', on_click=save_docente).props('unelevated color=primary w-full')

@ui.page('/archivio')
def archivio_page():
    if not app.storage.user.get('authenticated', False): 
        ui.navigate.to('/'); return

    state = {'search': '', 'date_start': None, 'date_end': None}
    table_ref = None

    # --- HELPER ---
    def format_date(dt_obj):
        """Formatta date o stringhe in DD/MM/YYYY"""
        if not dt_obj: return '-'
        try:
            if isinstance(dt_obj, str):
                return datetime.strptime(dt_obj, '%Y-%m-%d').strftime('%d/%m/%Y')
            return dt_obj.strftime('%d/%m/%Y')
        except: return str(dt_obj)

    # --- LOGICA EXPORT EXCEL/CSV ---
    async def export_excel():
        """Genera un file CSV (Excel compatibile) con i dati filtrati"""
        try:
            # 1. Recupera dati (con gli stessi filtri della tabella)
            rows = await asyncio.to_thread(AttestatiRepo.get_history, state['search'], state['date_start'], state['date_end'])
            
            if not rows:
                ui.notify("Nessun dato da esportare", color='warning')
                return

            # 2. Crea file temporaneo
            # 'utf-8-sig' è fondamentale per far leggere gli accenti a Excel
            # 'delete=False' perché NiceGUI deve poterlo leggere per inviarlo
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.csv', encoding='utf-8-sig', newline='') as tmp:
                # Usa ';' come delimitatore perché Excel italiano lo preferisce alla virgola
                writer = csv.writer(tmp, delimiter=';')
                
                # Intestazioni
                writer.writerow(['ID', 'Data Emissione', 'Corsista', 'Codice Fiscale', 'Corso', 'Scadenza', 'Stato'])
                
                today = date.today()
                
                # Scrittura righe
                for r in rows:
                    # Formattazione dati
                    d_em = format_date(r.get('DATA_EMISSIONE'))
                    d_sc = format_date(r.get('SCADENZA'))
                    
                    # Calcolo Stato (Testuale per Excel)
                    stato_str = "SCADUTO"
                    try:
                        scad_val = r.get('SCADENZA')
                        if scad_val:
                            if isinstance(scad_val, str):
                                scad_obj = datetime.strptime(scad_val, '%Y-%m-%d').date()
                            elif isinstance(scad_val, datetime):
                                scad_obj = scad_val.date()
                            else:
                                scad_obj = scad_val
                            
                            if scad_obj > today:
                                stato_str = "VALIDO"
                    except:
                        pass # Resta Scaduto se errore

                    writer.writerow([
                        r.get('ID', ''),
                        d_em,
                        r.get('CORSISTA', ''),
                        r.get('CF', ''),
                        r.get('CORSO', ''),
                        d_sc,
                        stato_str
                    ])
                
                tmp_path = tmp.name

            # 3. Avvia Download
            ui.download(tmp_path, 'Archivio_Attestati.csv')
            ui.notify("Export completato!", color='green')

        except Exception as e:
            ui.notify(f"Errore Export: {e}", color='red')

    # --- LOGICA TABELLA ---
    async def refresh_table():
        try:
            rows = await asyncio.to_thread(AttestatiRepo.get_history, state['search'], state['date_start'], state['date_end'])
            
            today = date.today()

            for r in rows:
                r['DATA_FMT'] = format_date(r.get('DATA_EMISSIONE'))
                r['SCADENZA_FMT'] = format_date(r.get('SCADENZA'))
                
                # Calcolo validità per Badge
                r['VALIDO'] = False 
                try:
                    scad_val = r.get('SCADENZA')
                    if scad_val:
                        if isinstance(scad_val, str):
                            scad = datetime.strptime(scad_val, '%Y-%m-%d').date()
                        elif isinstance(scad_val, datetime):
                            scad = scad_val.date()
                        else:
                            scad = scad_val 
                        
                        if scad and scad > today:
                            r['VALIDO'] = True
                except Exception as e:
                    print(f"Errore calcolo validità riga {r.get('ID')}: {e}")

            if table_ref: 
                table_ref.rows = rows
                table_ref.update()
        except Exception as e:
            ui.notify(f"Errore caricamento storico: {e}", color='red')

    # --- LAYOUT ---
    with ui.column().classes('w-full items-center p-8 max-w-screen-xl mx-auto bg-slate-50 min-h-screen'):
        
        with ui.row().classes('w-full items-center mb-6 justify-between'):
            with ui.row().classes('items-center gap-4'):
                ui.button(icon='arrow_back', on_click=lambda: ui.navigate.to('/dashboard')).props('flat round dense text-color=slate-700')
                ui.label('Archivio Attestati').classes('text-3xl font-bold text-slate-800')
                ui.icon('history', size='lg').classes('text-slate-400')

            # Bottone collegato alla funzione export_excel
            ui.button('Esporta Excel', icon='file_download', on_click=export_excel).props('outline color=green')

        with ui.card().classes('w-full p-4 mb-6 grid grid-cols-1 md:grid-cols-4 gap-4 items-end bg-white shadow-sm'):
            ui.input('Cerca...').props('outlined dense').classes('md:col-span-2').bind_value(state, 'search').on('keydown.enter', refresh_table)
            ui.input('Dal...').props('outlined dense type=date').bind_value(state, 'date_start')
            ui.button('Filtra', icon='filter_list', on_click=refresh_table).props('unelevated color=primary')

        cols = [
            {'name': 'DATA_FMT', 'label': 'Data', 'field': 'DATA_FMT', 'align': 'left', 'sortable': True},
            {'name': 'CORSISTA', 'label': 'Corsista', 'field': 'CORSISTA', 'align': 'left', 'sortable': True, 'classes': 'font-bold'},
            {'name': 'CF', 'label': 'CF', 'field': 'CF', 'align': 'left', 'classes': 'font-mono text-xs'},
            {'name': 'CORSO', 'label': 'Corso', 'field': 'CORSO', 'align': 'left'},
            {'name': 'SCADENZA_FMT', 'label': 'Scadenza', 'field': 'SCADENZA_FMT', 'align': 'center'},
            {'name': 'status', 'label': 'Stato', 'field': 'status', 'align': 'center'},
        ]
        
        table_ref = ui.table(columns=cols, rows=[], row_key='ID').classes('w-full shadow-md bg-white')

        table_ref.add_slot('body-cell-status', r'''
            <q-td key="status" :props="props">
                <q-badge :color="props.row.VALIDO ? 'green' : 'red'" 
                         :label="props.row.VALIDO ? 'VALIDO' : 'SCADUTO'" />
            </q-td>
        ''')
        
        ui.timer(0.1, refresh_table, once=True)

@ui.page('/scadenzario')
def scadenzario_page():
    if not app.storage.user.get('authenticated', False): 
        ui.navigate.to('/'); return

    # Stato
    state = {
        'days_lookahead': 60, 
        'search': '',
        'filter_mode': 'in_scadenza',
        # --- STATO PER LA MAIL ---
        'mail_to': '',
        'mail_subject': '',
        'mail_body': ''
    }
    table_ref = None
    email_dialog = None # Riferimento al popup

    # --- HELPER ---
    def format_date(dt_obj):
        if not dt_obj: return '-'
        try:
            if isinstance(dt_obj, str): 
                dt_obj = datetime.strptime(dt_obj, '%Y-%m-%d')
            return dt_obj.strftime('%d/%m/%Y')
        except: return str(dt_obj)

    # --- LOGICA GESTIONE MAIL ---
    def open_email_dialog(e):
        """Apre la finestra per comporre la mail con dati precompilati"""
        row = e.args
        corsista = row.get('CORSISTA', 'Dipendente') # Contiene già "Cognome Nome"
        corso = row.get('CORSO', 'Corso')
        scadenza = row.get('SCADENZA_FMT', 'Data ignota')
        
        # Recuperiamo il nome Ente se disponibile nella riga, altrimenti generico
        ente_nome = row.get('ENTE', 'Spett.le Azienda')
        
        # 1. Precompila i campi
        state['mail_to'] = '' 
        state['mail_subject'] = f"Rinnovo Corso: {corso} - {corsista}"
        
        # Template del messaggio AGGIORNATO
        state['mail_body'] = (
            f"Gentile {ente_nome},\n\n"
            f"Ti ricordiamo che l'attestato di {corsista} per il corso di formazione '{corso}' "
            f"risulta in scadenza il {scadenza}.\n\n"
            "Per mantenere la validità della certificazione, è necessario programmare il corso di aggiornamento.\n"
            "Restiamo a disposizione per organizzare le date.\n\n"
            "Cordiali saluti,\n"
            "Segreteria Formazione"
        )
        
        # 2. Apre il popup
        email_dialog.open()

    async def send_email_action():
        """
        Funzione che invia effettivamente la mail tramite SMTP.
        """
        destinatario = state['mail_to']
        oggetto = state['mail_subject']
        corpo = state['mail_body']

        if not destinatario:
            ui.notify("Inserisci un indirizzo email destinatario!", type='warning')
            return

        # --- CONFIGURAZIONE MITTENTE (DA COMPILARE) ---
        # Se usi Gmail, devi generare una "App Password" nelle impostazioni di sicurezza Google
        SMTP_SERVER = "smtp.gmail.com"
        SMTP_PORT = 587
        SENDER_EMAIL = "lorenzo.ricci.onaws@gmail.com" # <--- INSERISCI QUI LA TUA EMAIL
        SENDER_PASSWORD = "Occhiali28!" # <--- INSERISCI QUI LA TUA PASSWORD
        
        # Feedback visivo
        ui.notify("Connessione al server di posta...", spinner=True)
        
        try:
            # Eseguiamo l'invio in un thread separato per non bloccare l'interfaccia
            def invia_reale():
                msg = MIMEMultipart()
                msg['From'] = SENDER_EMAIL
                msg['To'] = destinatario
                msg['Subject'] = oggetto
                msg.attach(MIMEText(corpo, 'plain'))

                server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
                server.starttls() # Sicurezza TLS
                server.login(SENDER_EMAIL, SENDER_PASSWORD)
                text = msg.as_string()
                server.sendmail(SENDER_EMAIL, destinatario, text)
                server.quit()

            await asyncio.to_thread(invia_reale)
            
            ui.notify(f"Email inviata correttamente a {destinatario}!", type='positive')
            email_dialog.close()

        except Exception as e:
            ui.notify(f"Errore invio: {e}", color='red', multi_line=True)

    # --- LOGICA RECUPERO DATI ---
    async def refresh_table():
        try:
            all_rows = await asyncio.to_thread(AttestatiRepo.get_history, state['search'])
            
            if not all_rows:
                if table_ref: table_ref.update()
                return

            today = date.today()
            limit_date = today + timedelta(days=int(state['days_lookahead']))
            
            filtered_rows = []
            
            for r in all_rows:
                scad_obj = None
                try:
                    raw = r.get('SCADENZA')
                    if raw:
                        if isinstance(raw, str): 
                            scad_obj = datetime.strptime(raw, '%Y-%m-%d').date()
                        elif isinstance(raw, datetime):
                            scad_obj = raw.date()
                        elif isinstance(raw, date): 
                            scad_obj = raw
                except: pass

                if not scad_obj: continue 

                include = False
                days_left = (scad_obj - today).days
                
                if state['filter_mode'] == 'scaduti':
                    if days_left < 0: include = True
                elif state['filter_mode'] == 'in_scadenza':
                    if -60 <= days_left <= state['days_lookahead']: include = True
                else: 
                    include = True

                if include:
                    r['GIORNI_RIMASTI'] = days_left
                    r['SCADENZA_FMT'] = format_date(scad_obj)
                    
                    if days_left < 0:
                        r['STATUS_COLOR'] = 'red'
                        r['STATUS_LABEL'] = f'SCADUTO da {abs(days_left)} gg'
                    elif days_left <= 30:
                        r['STATUS_COLOR'] = 'orange'
                        r['STATUS_LABEL'] = f'Scade tra {days_left} gg'
                    else:
                        r['STATUS_COLOR'] = 'green'
                        r['STATUS_LABEL'] = f'Scade tra {days_left} gg'
                    
                    filtered_rows.append(r)

            filtered_rows.sort(key=lambda x: x['GIORNI_RIMASTI'])

            if table_ref: 
                table_ref.rows = filtered_rows
                table_ref.update()
                
        except Exception as e:
            ui.notify(f"Errore caricamento scadenze: {e}", color='red')

    # --- LAYOUT ---
    with ui.column().classes('w-full items-center p-8 bg-slate-50 min-h-screen'):
        
        # Header
        with ui.row().classes('w-full items-center mb-6 justify-between max-w-screen-xl'):
            with ui.row().classes('items-center gap-4'):
                ui.button(icon='arrow_back', on_click=lambda: ui.navigate.to('/dashboard')).props('flat round dense text-color=slate-700')
                ui.label('Scadenzario Corsi').classes('text-3xl font-bold text-slate-800')
                ui.icon('alarm', size='lg').classes('text-orange-500')
            
            with ui.row().classes('gap-4'):
                ui.badge('Controlla regolarmente!', color='orange').props('outline')

        # Barra Strumenti
        with ui.card().classes('w-full p-4 mb-6 flex flex-wrap items-center gap-4 bg-white shadow-sm max-w-screen-xl'):
            with ui.button_group().props('unelevated'):
                ui.button('In Scadenza', on_click=lambda: (state.update({'filter_mode': 'in_scadenza'}), refresh_table())).props('color=orange')
                ui.button('Già Scaduti', on_click=lambda: (state.update({'filter_mode': 'scaduti'}), refresh_table())).props('color=red')
                ui.button('Tutto', on_click=lambda: (state.update({'filter_mode': 'tutti'}), refresh_table())).props('color=grey')

            ui.separator().props('vertical')

            with ui.row().classes('items-center gap-2'):
                ui.label('Orizzonte temporale:').classes('text-sm text-gray-600')
                slider = ui.slider(min=30, max=365, step=30, value=60).props('label-always color=primary').classes('w-48')
                slider.bind_value(state, 'days_lookahead')
                slider.on('change', refresh_table)
                ui.label('giorni').classes('text-sm')

            ui.separator().props('vertical')
            ui.input('Cerca Azienda o Corsista').props('outlined dense').classes('flex-grow').bind_value(state, 'search').on('keydown.enter', refresh_table)
            ui.button(icon='refresh', on_click=refresh_table).props('flat round')

        # Tabella
        cols = [
            {'name': 'SCADENZA_FMT', 'label': 'Data Scadenza', 'field': 'SCADENZA_FMT', 'align': 'left', 'sortable': True},
            {'name': 'status', 'label': 'Stato', 'field': 'status', 'align': 'left'},
            {'name': 'CORSISTA', 'label': 'Corsista', 'field': 'CORSISTA', 'align': 'left', 'sortable': True, 'classes': 'font-bold'},
            {'name': 'CORSO', 'label': 'Corso da Rinnovare', 'field': 'CORSO', 'align': 'left'},
            {'name': 'CF', 'label': 'CF', 'field': 'CF', 'align': 'left', 'classes': 'text-xs text-gray-500'},
            {'name': 'azioni', 'label': '', 'field': 'azioni', 'align': 'right'},
        ]
        
        table_ref = ui.table(columns=cols, rows=[], row_key='ID').classes('w-full shadow-md bg-white max-w-screen-xl')

        table_ref.add_slot('body-cell-status', r'''
            <q-td key="status" :props="props">
                <q-chip :color="props.row.STATUS_COLOR" text-color="white" dense>
                    {{ props.row.STATUS_LABEL }}
                </q-chip>
            </q-td>
        ''')

        table_ref.add_slot('body-cell-azioni', r'''
            <q-td key="azioni" :props="props">
                <q-btn icon="mail" label="Avvisa" size="sm" flat color="primary" @click="$parent.$emit('email', props.row)" />
            </q-td>
        ''')
        
        # Collegamento evento apertura dialog
        table_ref.on('email', open_email_dialog)
        
        ui.timer(0.1, refresh_table, once=True)

    # --- DIALOG INVIO EMAIL (NUOVO) ---
    with ui.dialog() as email_dialog, ui.card().classes('w-full max-w-2xl p-6 gap-4'):
        
        # Intestazione
        with ui.row().classes('w-full items-center justify-between mb-2'):
            with ui.row().classes('items-center gap-2'):
                ui.icon('send', size='md').classes('text-primary')
                ui.label('Invia Avviso Scadenza').classes('text-xl font-bold text-slate-800')
            ui.button(icon='close', on_click=email_dialog.close).props('flat round dense color=grey')

        # Form
        ui.input('Destinatario (Email)').bind_value(state, 'mail_to').props('outlined dense type=email').classes('w-full')
        ui.input('Oggetto').bind_value(state, 'mail_subject').props('outlined dense').classes('w-full font-bold')
        
        ui.label('Testo del messaggio:').classes('text-sm text-gray-500 mt-2')
        ui.textarea().bind_value(state, 'mail_body').props('outlined dense rows=10').classes('w-full')

        # Footer
        with ui.row().classes('w-full justify-end mt-4 gap-2'):
            ui.button('Annulla', on_click=email_dialog.close).props('flat color=grey')
            ui.button('Invia Email', on_click=send_email_action).props('unelevated color=primary icon=send')

if __name__ in {"__main__", "__mp_main__"}:
    ui.run(title="WorkSafeManager", storage_secret='secret_key', reload=True ,port=8001)